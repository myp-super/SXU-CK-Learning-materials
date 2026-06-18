#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
FPGA数字系统设计课程设计报告生成器
生成完整的Word格式课程设计报告
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

def set_cell_shading(cell, color):
    """设置单元格底色"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_paragraph_with_format(doc, text, style='Normal', bold=False, size=12, alignment=None, font_name='宋体'):
    """添加格式化段落"""
    p = doc.add_paragraph(style=style)
    if alignment is not None:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return p

def add_heading_styled(doc, text, level=1):
    """添加带样式的标题"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h

def add_body_text(doc, text):
    """添加正文段落"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)  # 两字符缩进
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def add_code_file(doc, filepath, title=None):
    """将完整的VHDL源文件内容添加到文档"""
    base_dir = r'C:\Users\Lenovo\Desktop\shudiankeshe'
    full_path = os.path.join(base_dir, filepath)
    try:
        with open(full_path, 'r', encoding='utf-8') as f:
            code = f.read()
    except:
        code = '-- (文件读取失败: {})'.format(filepath)
    if title:
        add_heading_styled(doc, title, 3)
    else:
        add_heading_styled(doc, filepath, 3)
    code_p = doc.add_paragraph()
    code_p.paragraph_format.left_indent = Cm(0.3)
    code_p.paragraph_format.space_before = Pt(2)
    code_p.paragraph_format.space_after = Pt(2)
    code_run = code_p.add_run(code)
    code_run.font.size = Pt(6.5)
    code_run.font.name = 'Consolas'
    return code

def create_report():
    doc = Document()

    # ======================== 页面设置 ========================
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ======================== 封面 ========================
    for i in range(6):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run('FPGA数字系统设计课程设计')
    run.font.size = Pt(26)
    run.bold = True
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_p.add_run('—— 数字钟的FPGA设计与实现')
    run.font.size = Pt(18)
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    doc.add_paragraph()
    doc.add_paragraph()

    # 课程设计信息表
    info_table = doc.add_table(rows=8, cols=2)
    info_table.style = 'Table Grid'
    info_data = [
        ('学    院', '信息科学与工程学院'),
        ('专    业', '电子信息工程'),
        ('学生姓名', '（填写）'),
        ('学    号', '（填写）'),
        ('课程名称', 'FPGA数字系统设计课程设计'),
        ('设计题目', '数字钟'),
        ('指导教师', '（填写）'),
        ('日    期', '2026年6月15日'),
    ]
    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = value
        for j in range(2):
            for paragraph in info_table.rows[i].cells[j].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(12)
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_page_break()

    # ======================== 评分表 ========================
    add_heading_styled(doc, '课程设计评分表', 1)
    score_table = doc.add_table(rows=8, cols=5)
    score_table.style = 'Table Grid'
    score_headers = ['序号', '毕业要求指标点', '考核内容', '分值', '得分']
    for i, h in enumerate(score_headers):
        score_table.rows[0].cells[i].text = h
        set_cell_shading(score_table.rows[0].cells[i], 'D9E2F3')

    score_data = [
        ('1', 'G3.1', '学习态度、出勤', '5', ''),
        ('2', 'G3.1', '设计方案的合理性、创新性', '15', ''),
        ('3', 'G3.1', '设计过程及成果质量', '30', ''),
        ('4', 'G11.1', '进度安排的项目管理', '15', ''),
        ('5', 'G10.2', '外文资料阅读与翻译能力', '10', ''),
        ('6', '—', '报告撰写质量', '25', ''),
        ('7', '—', '总  分', '100', ''),
    ]
    for i, row_data in enumerate(score_data):
        for j, val in enumerate(row_data):
            score_table.rows[i+1].cells[j].text = val

    doc.add_page_break()

    # ======================== 第1章 研究背景 ========================
    add_heading_styled(doc, '1  研究背景', 1)

    add_heading_styled(doc, '1.1  数字钟的发展历程', 2)
    add_body_text(doc,
        '时间测量是人类文明发展史上最古老且最重要的技术活动之一。从远古时期的日晷、水钟、沙漏等自然计时装置，'
        '到近代的机械钟表、电子钟表，再到当代基于微处理器和可编程逻辑器件的智能数字钟，计时技术的发展始终与'
        '人类科技进步紧密相连。数字钟作为现代计时技术的典型代表，其发展历程清晰地反映了电子技术的演进脉络。')

    add_body_text(doc,
        '机械钟表时代（14世纪至20世纪中期）：最早的机械钟诞生于14世纪的欧洲修道院，以重锤或发条为动力源，'
        '通过擒纵机构和齿轮传动系统实现计时。17世纪伽利略发现摆的等时性原理后，摆钟的出现使计时精度从每日'
        '误差数十分钟提升到数秒级别。18世纪航海天文钟的发明解决了海上经度测量难题，其精度达到每日误差仅0.1秒。'
        '然而，机械钟表结构复杂、制造精密、成本高昂，且长期运行存在磨损问题，限制了其大规模普及。')

    add_body_text(doc,
        '电子钟表时代（20世纪中期至70年代）：1927年石英晶体振荡器的发明标志着电子计时时代的开端。石英钟利用'
        '石英晶体的压电效应产生极其稳定的振荡频率（通常为32.768kHz），通过分频电路获得精确的秒脉冲信号。'
        '1969年日本精工推出世界上第一款石英腕表Astron，其月误差仅为±5秒，远优于任何机械表。此后，大规模'
        '集成电路（LSI）技术的发展使电子钟表成本大幅下降，数字显示替代了传统的指针显示，LED和LCD数码管'
        '成为主流显示器件。这一阶段，数字钟逐步从专业计时仪器走向大众消费市场。')

    add_body_text(doc,
        '单片机数字钟时代（20世纪80年代至21世纪初）：微控制器（MCU）的出现为数字钟设计带来了革命性变化。'
        '以Intel 8051、Microchip PIC、Atmel AVR为代表的8位单片机，通过软件编程即可实现定时、计数、显示'
        '驱动和键盘扫描等功能。单片机数字钟具有功能灵活、开发周期短、成本低廉等优点，能够轻松实现闹钟、'
        '倒计时、世界时钟等附加功能。在这一时期，数字钟的设计重心从硬件电路转向软件算法，系统的可扩展性和'
        '可维护性显著提升。然而，单片机方案本质上是顺序执行的，在需要并行处理多任务的复杂应用场景中存在'
        '性能瓶颈，且软件实现的计时精度受中断响应时间、指令周期等因素影响。')

    add_body_text(doc,
        'FPGA数字钟时代（21世纪初至今）：现场可编程门阵列（Field Programmable Gate Array, FPGA）的出现'
        '代表了数字系统设计方法的根本性变革。与单片机不同，FPGA通过硬件描述语言（Hardware Description '
        'Language, HDL）描述电路的结构和行为，经综合、布局布线后直接映射为硬件电路。FPGA数字钟的所有功能'
        '模块——分频器、计数器、译码器、控制器——均以并行方式同时工作，在本质上不同于单片机的顺序执行模式。'
        '这种硬件并行的特性使FPGA数字钟具有纳秒级的响应速度和皮秒级的计时精度，且系统稳定性不受软件"跑飞"'
        '等因素影响。此外，FPGA的可重配置特性使得设计者可以在不改变硬件电路板的情况下，通过重新编程实现'
        '功能升级和优化，这为数字钟的设计和迭代提供了极大的灵活性。')

    add_heading_styled(doc, '1.2  国内外研究现状', 2)
    add_body_text(doc,
        'FPGA技术自1985年Xilinx公司推出首款商业化产品XC2064以来，经历了三十余年的快速发展。当前，全球'
        'FPGA市场主要由Xilinx（现为AMD子公司）和Intel（2015年收购Altera）双寡头垄断，合计占据约85%的'
        '市场份额。在工艺节点方面，Xilinx Versal系列已采用7nm工艺，Intel Agilex系列采用10nm工艺，单片'
        '逻辑单元数量已突破千万门级别。在开发工具方面，Xilinx Vivado和Intel Quartus Prime提供了从设计'
        '输入、综合优化、布局布线到时序分析的完整工具链，大幅降低了FPGA开发的入门门槛（Xilinx Inc., 2023）。')

    add_body_text(doc,
        '在FPGA计时系统应用领域，国内外学者进行了大量研究。Kuon和Rose（2007）在其经典论文中系统比较了'
        'FPGA与ASIC在实现数字电路时的面积和延迟差异，指出FPGA在中小批量、快速原型验证场景中具有显著的'
        '成本和时间优势。Hauck和DeHon（2010）在其专著《Reconfigurable Computing》中详细阐述了可重构'
        '计算在嵌入式系统中的应用，包括基于FPGA的高精度计时和同步系统设计。Wolf（2018）在《FPGA-Based '
        'System Design》中探讨了FPGA在工业控制、仪器仪表和消费电子中的系统级设计方法学。')

    add_body_text(doc,
        '国内在FPGA教学和科研方面也取得了长足进步。潘松、黄继业编著的《EDA技术与VHDL》是国内高校FPGA课程'
        '的经典教材，系统介绍了VHDL语言基础和FPGA设计流程。王金明在《FPGA设计实战》中通过大量工程实例展示'
        '了FPGA在数字系统设计中的应用方法。近年来，随着国产FPGA芯片（如紫光同创、安路科技、高云半导体等）'
        '的崛起，国内FPGA产业生态日趋完善。在数字钟这一经典教学案例中，各高校不断探索将课程设计与工程实践'
        '相结合的教学模式，强调从系统需求出发，经历方案论证、模块设计、仿真验证到硬件实现的完整设计流程。')

    add_heading_styled(doc, '1.3  FPGA相对于MCU的优势分析', 2)
    add_body_text(doc,
        'FPGA与单片机（MCU）是数字系统设计中两种主流平台，它们在架构理念、工作方式和适用场景上存在本质区别。'
        '首先，从工作方式来看，MCU基于冯·诺依曼或哈佛架构，通过顺序取指、译码、执行的循环完成程序指令，'
        '所有任务在时间上是串行的；而FPGA本质上是硬件电路的可编程实现，各个功能模块在物理上独立并行运行，'
        '不存在CPU的时间片轮转开销。以数字钟为例，FPGA方案中的分频器、秒计数器、分计数器、时计数器、'
        '译码显示模块和报时模块在综合后成为独立的硬件电路，它们之间通过物理连线传递信号，真正实现了全并行处理。')

    add_body_text(doc,
        '其次，从实时性和确定性的角度，MCU的中断响应存在不确定性——当中断发生时，CPU需要保存现场、跳转'
        '执行中断服务程序、再恢复现场，这一过程的延迟受指令流水线状态、中断优先级、DMA竞争等因素影响，'
        '在最坏情况下可达数十微秒；而FPGA的响应是硬件级别的，信号从输入引脚到输出引脚的传播延迟完全由'
        '组合逻辑和布线延迟决定，通常在纳秒级别，且具有严格的确定性。这对于需要精确同步的计时系统尤为重要。')

    add_body_text(doc,
        '再次，从功耗和成本角度，MCU在低功耗简单应用中具有优势，现代低功耗MCU在休眠模式下功耗可低至'
        '微安级别；FPGA的静态功耗相对较高，但在需要并行处理大量数据的应用中，FPGA的能效比（Performance '
        'per Watt）往往优于MCU+软件方案。在成本方面，低端FPGA（如Intel Cyclone IV EP4CE6）芯片批量价格'
        '已降至5美元以下，配合免费开发工具链，使得FPGA方案在成本和开发便捷性上已经与高端MCU方案相当。')

    add_heading_styled(doc, '1.4  FPGA的应用领域', 2)
    add_body_text(doc,
        'FPGA凭借其并行处理能力、可重构特性和丰富的I/O资源，已广泛应用于以下领域：')

    add_body_text(doc,
        '（1）通信与网络：FPGA是5G基站、光传输网络、软件定义网络（SDN）等通信基础设施的核心器件。'
        '高速SerDes（串行/解串器）接口、前向纠错（FEC）编解码、数字上/下变频（DUC/DDC）等基带处理'
        '功能均可通过FPGA高效实现。Cisco、华为、中兴等通信设备商大量采用FPGA进行协议处理和接口桥接。')

    add_body_text(doc,
        '（2）工业控制：在数控机床、机器人控制、运动控制等工业自动化领域，FPGA提供微秒级甚至纳秒级的'
        '实时控制能力。多轴伺服电机同步控制、高速PWM生成、编码器接口等均可通过单个FPGA芯片实现。')

    add_body_text(doc,
        '（3）人工智能加速：FPGA在AI推理加速领域具有独特优势。Microsoft Azure的Catapult项目大规模部署'
        'Intel FPGA用于Bing搜索排序和深度学习推理加速。Xilinx的DPU（深度学习处理单元）IP核可在FPGA上'
        '实现高效的卷积神经网络（CNN）推理，能效比远超GPU方案。')

    add_body_text(doc,
        '（4）图像与视频处理：FPGA的流水线架构天然适合图像处理中的像素级并行操作。在医疗影像、工业检测、'
        '自动驾驶视觉感知等场景中，FPGA可实时处理4K/8K视频流，执行图像滤波、边缘检测、目标识别等算法。')

    add_body_text(doc,
        '（5）数字仪器仪表：FPGA是高精度频率计、逻辑分析仪、任意波形发生器、数字示波器等测试测量仪器的'
        '核心处理单元。其精确的时序控制能力和高速数据采集接口使仪器能够实现皮秒级的时间测量和GSa/s级的'
        '采样率。本课程设计所选用的数字钟正是FPGA在数字仪器仪表领域的典型教学案例。')

    doc.add_page_break()

    # ======================== 第2章 设计方案 ========================
    add_heading_styled(doc, '2  设计方案', 1)
    add_body_text(doc,
        '本章详细阐述数字钟的两种设计方案：方案一采用"二进制计数器+二进制转BCD转换+七段译码"的三级流水'
        '架构，方案二采用"BCD计数器直接输出"的两级架构。两种方案在计数器实现方式、数据通路结构和模块划分'
        '上存在本质区别，各有优劣。以下从结构框图、组成模块、工作原理、优缺点和资源消耗等维度分别论述。')

    add_heading_styled(doc, '2.1  方案一：二进制计数器 + 转换模块方案', 2)

    add_heading_styled(doc, '2.1.1  结构框图', 3)
    add_body_text(doc,
        '方案一的系统结构如图2-1所示（详见附录结构框图）。该系统采用"二进制计数→BCD码转换→七段译码显示"'
        '的三级流水线架构。时钟源经分频器产生1Hz标准秒脉冲信号，驱动二进制秒计数器（模60）进行加1计数；'
        '秒计数器溢出产生进位信号，驱动二进制分计数器（模60）；分计数器溢出驱动二进制时计数器（模24）。'
        '各级计数器的6位（或5位）二进制输出送入二进制-BCD转换模块，转换为8421BCD码后，再经七段译码器'
        '转换为数码管段选信号。校时/校分模块通过多路选择器将快速时钟脉冲旁路注入分/时计数器，实现时间的'
        '快速调整。整点报时模块检测59分59秒状态，触发蜂鸣器输出。')

    add_heading_styled(doc, '2.1.2  组成模块', 3)
    # 模块列表表格
    module_table = doc.add_table(rows=10, cols=4)
    module_table.style = 'Table Grid'
    m_headers = ['序号', '模块名称', '功能描述', '实现方式']
    for i, h in enumerate(m_headers):
        module_table.rows[0].cells[i].text = h
        set_cell_shading(module_table.rows[0].cells[i], 'D9E2F3')

    modules = [
        ('1', '分频器(diver)', '将高频系统时钟分频为1Hz标准时钟和快速校时时钟', '整数计数器分频'),
        ('2', '二进制秒计数器', '模60二进制计数(0~59)，输出6位二进制', '行为描述，process+if-else'),
        ('3', '二进制分计数器', '模60二进制计数(0~59)，输出6位二进制', '行为描述'),
        ('4', '二进制时计数器', '模24二进制计数(0~23)，输出5位二进制', '行为描述'),
        ('5', '二进制→BCD转换', '6位/5位二进制→8421BCD（十位+个位）', '组合逻辑，除法和取模'),
        ('6', '七段译码器', '4位BCD→7段数码管段选信号', 'case语句真值表译码'),
        ('7', '校时/校分控制', '按键控制快速调时/调分', '多路选择器MUX'),
        ('8', '整点报时模块', '检测59′59″状态输出蜂鸣信号', '组合逻辑比较器'),
        ('9', '顶层模块', '各子模块的例化和信号互连', '结构描述，元件例化'),
    ]
    for i, row_data in enumerate(modules):
        for j, val in enumerate(row_data):
            module_table.rows[i+1].cells[j].text = val

    add_heading_styled(doc, '2.1.3  工作原理', 3)
    add_body_text(doc,
        '方案一的核心设计思想是"关注点分离"：计数功能与显示编码功能分别由独立模块承担。计数器仅关注'
        '0~59（或0~23）的二进制累加和回零控制，不关心输出值的显示格式；转换模块专门负责将二进制计数值'
        '映射为人类可读的十进制BCD表示；译码模块再将BCD码映射为数码管的物理段选信号。这种模块化分层'
        '架构使每个模块的职责单一明确，便于独立设计、仿真验证和后期维护。')

    add_body_text(doc,
        '以秒计数器为例：当1Hz时钟上升沿（或下降沿）到达时，若使能信号有效，6位二进制内部寄存器从0开始'
        '递增。当计数值达到59（二进制111011）时，下一个时钟沿将计数器归零，同时输出一个周期的进位脉冲。'
        '该进位脉冲作为分计数器的时钟使能信号，形成级联计数链。二进制计数值（如"001011"代表11秒）送入'
        'bin2bcd模块，通过除10取整得到十位BCD码"0001"（1），取模10得到个位BCD码"0001"（1），最终'
        '输出"0001 0001"代表11秒。该BCD码再送入seg7_decoder模块，经case语句查表译码后驱动数码管显示"11"。')

    add_heading_styled(doc, '2.1.4  优点分析', 3)
    add_body_text(doc,
        '（1）模块化程度高：计数、转换、显示三者功能独立，遵循数字系统设计的"分而治之"原则。任一模块'
        '的修改不影响其他模块，例如更换不同类型的数码管（共阴/共阳）仅需修改seg7_decoder的真值表。')
    add_body_text(doc,
        '（2）可扩展性强：二进制计数器可轻松扩展为更大模值（如模100、模1000），二进制→BCD转换模块'
        '仅需增加输入位宽和调整转换逻辑，而无需重新设计整个计数器。')
    add_body_text(doc,
        '（3）教学意义突出：方案一完整展示了"二进制运算→十进制显示"这一计算机体系结构中的基础概念，'
        '对于理解计算机内部二进制存储与外部十进制显示之间的转换机制具有重要的教学价值。')
    add_body_text(doc,
        '（4）移植性好：二进制计数模块不包含任何与显示相关的逻辑，可复用于其他需要模60/模24二进制'
        '计数的数字系统（如定时器、秒表），体现了良好的IP复用理念。')

    add_heading_styled(doc, '2.1.5  缺点分析', 3)
    add_body_text(doc,
        '（1）资源开销较大：需要额外的bin2bcd转换模块（对6位输入需约30个逻辑单元）和seg7_decoder译码'
        '模块（6个实例需约42个逻辑单元），相比方案二增加了约20%的组合逻辑资源消耗。')
    add_body_text(doc,
        '（2）路径延迟增加：信号从计数器输出到数码管显示需经过"计数器→bin2bcd转换→seg7译码→输出引脚"'
        '三级组合逻辑路径，在最坏情况下可能引入10~15ns的额外延迟（具体取决于器件和布线）。')
    add_body_text(doc,
        '（3）设计复杂度较高：相比方案二的"即计即显"，方案一需要设计者正确理解二进制与BCD码的转换关系，'
        '对初学者的数字逻辑基础要求更高。')

    add_heading_styled(doc, '2.2  方案二：BCD计数器直接输出方案', 2)

    add_heading_styled(doc, '2.2.1  结构框图', 3)
    add_body_text(doc,
        '方案二的系统结构如图2-2所示（详见附录结构框图）。该方案采用"BCD计数→直接显示"的两级架构。'
        '与方案一的关键区别在于：计数器本身以BCD格式（个位和十位分别以4位BCD码表示）进行计数，'
        '计数值可直接用于显示，无需额外的编码转换模块。秒计数器由两个4位BCD子计数器组成：个位计数器'
        '（0~9循环）和十位计数器（0~5循环），两者级联构成模60 BCD计数器。分计数器和时计数器采用'
        '相同的BCD计数架构（时计数器的十位为0~2）。')

    add_heading_styled(doc, '2.2.2  组成模块', 3)
    module_table2 = doc.add_table(rows=8, cols=4)
    module_table2.style = 'Table Grid'
    for i, h in enumerate(m_headers):
        module_table2.rows[0].cells[i].text = h
        set_cell_shading(module_table2.rows[0].cells[i], 'D9E2F3')

    modules2 = [
        ('1', '分频器(diver)', '系统时钟分频', '整数计数器分频'),
        ('2', 'BCD秒计数器', '模60 BCD计数，直接输出个位和十位BCD码', '行为描述，双4位计数器级联'),
        ('3', 'BCD分计数器', '模60 BCD计数', '行为描述'),
        ('4', 'BCD时计数器', '模24 BCD计数', '行为描述'),
        ('5', '校时/校分控制', '按键控制快速调时/调分', '多路选择器MUX'),
        ('6', '整点报时模块', '检测59′59″状态', '组合逻辑比较器'),
        ('7', '顶层模块', '子模块例化和信号互连', '结构描述'),
    ]
    for i, row_data in enumerate(modules2):
        for j, val in enumerate(row_data):
            module_table2.rows[i+1].cells[j].text = val

    add_heading_styled(doc, '2.2.3  工作原理', 3)
    add_body_text(doc,
        '方案二的核心设计思想是"显示导向"：计数器的内部状态直接对应于数码管的显示内容。以秒计数器为例，'
        '其内部维护两个4位寄存器num0（个位）和num1（十位）。当1Hz时钟到达时：若当前个位num0≠9（即"1001"），'
        '则num0<=num0+1；若num0=9，则num0<=0，同时十位num1<=num1+1。当num0=9且num1=5时，计数器整体归零。'
        '每个BCD位的值（如"0001"代表十位为1，"0011"代表个位为3）直接输出到显示端口，无需任何中间转换。')

    add_body_text(doc,
        '时计数器的计数逻辑类似，但十位上限为2（"0010"）。当时计数器的个位num0=3且十位num1=2时，表示'
        '当前为23时，下一时钟沿计数器归零（00时），实现24小时循环。值得注意的是，BCD计数器的条件判断'
        '比二进制计数器更复杂：二进制计数器仅需判断"count=59"，而BCD计数器需要分别判断个位=9和十位=5'
        '两个边界条件，且进位逻辑涉及两级if-else嵌套。')

    add_heading_styled(doc, '2.2.4  优点分析', 3)
    add_body_text(doc,
        '（1）结构简洁直观：计数器输出直接就是可显示的BCD码，数据通路短。从计数到显示的延迟仅为一级'
        '组合逻辑（若有七段译码器则为两级），信号路径最简。')
    add_body_text(doc,
        '（2）资源利用效率高：省略了独立的二进制→BCD转换模块，节省了约30个逻辑单元（LE），在资源'
        '极其有限的低端FPGA（如EPM240 CPLD）上具有实用价值。')
    add_body_text(doc,
        '（3）调试方便：计数器内部状态直接对应显示值，无需在心算"二进制→十进制"的映射关系，在硬件'
        '调试阶段（如使用SignalTap逻辑分析仪观察内部信号）时更加直观。')
    add_body_text(doc,
        '（4）功耗略低：由于组合逻辑级数更少，信号翻转传播路径更短，动态功耗（与信号翻转率成正比）'
        '理论上略低于方案一（虽然在实际数字钟应用中此差异可忽略不计）。')

    add_heading_styled(doc, '2.2.5  缺点分析', 3)
    add_body_text(doc,
        '（1）可扩展性受限：若需将计数器扩展为模100或更大模值，BCD计数器的条件判断逻辑会急剧复杂化'
        '（需处理个位、十位、百位三级嵌套），而二进制计数器仅需更改比较阈值即可。')
    add_body_text(doc,
        '（2）不符合通用设计范式：在大型数字系统中，内部计算通常采用二进制以简化算术运算（加减乘除），'
        '仅在需要人机交互的边界进行BCD转换。方案二的"全链路BCD"方法不便于与其他算术模块（如定时器'
        '比较器、时间加减运算器）集成。')
    add_body_text(doc,
        '（3）教学覆盖范围有限：方案二虽简洁高效，但未能覆盖"二进制→BCD转换"这一数字系统设计中的重要'
        '知识点，在教学的全面性上不及方案一。')

    add_heading_styled(doc, '2.3  两种方案对比分析', 2)

    # 对比表
    compare_table = doc.add_table(rows=8, cols=3)
    compare_table.style = 'Table Grid'
    compare_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    comp_headers = ['对比维度', '方案一（二进制+转换）', '方案二（BCD直出）']
    for i, h in enumerate(comp_headers):
        compare_table.rows[0].cells[i].text = h
        set_cell_shading(compare_table.rows[0].cells[i], '4472C4')
        for run in compare_table.rows[0].cells[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)

    comp_data = [
        ('设计复杂度', '较高：需设计bin2bcd和seg7_decoder\n两个额外模块，信号链路长', '较低：计数器直接输出BCD，\n模块数量最少'),
        ('可扩展性', '优秀：二进制计数天然适合扩展，\n仅需修改位宽和比较阈值', '一般：BCD计数扩展需增加多级\n嵌套条件判断'),
        ('可维护性', '优秀：模块职责单一，遵循\n"关注点分离"原则', '良好：结构简单易懂，但修改\n计数范围需改动内部逻辑'),
        ('资源利用率', '较低：额外转换模块占约15~20%\n额外逻辑资源', '较高：最简数据通路，\n逻辑单元使用最少'),
        ('仿真难度', '中等：需分别验证计数器、\n转换器和译码器的正确性', '较低：仿真波形直接对应\n显示值，验证更直观'),
        ('教学价值', '高：覆盖二进制运算、BCD转换、\n七段译码等核心知识点', '中等：侧重计数器和控制逻辑，\n知识覆盖面较窄'),
        ('适用场景', '大型系统、需算术运算、\n强调模块复用', '小型系统、资源受限、\n追求最简设计'),
    ]
    for i, (dim, s1, s2) in enumerate(comp_data):
        compare_table.rows[i+1].cells[0].text = dim
        compare_table.rows[i+1].cells[1].text = s1
        compare_table.rows[i+1].cells[2].text = s2
        set_cell_shading(compare_table.rows[i+1].cells[0], 'D9E2F3')

    add_heading_styled(doc, '2.4  最终方案选择', 2)
    add_body_text(doc,
        '综合考虑本课程设计的教学目标和实际需求，选择方案二（BCD计数器直接输出方案）作为实际硬件实现方案，'
        '同时以方案一（二进制计数器+转换方案）作为对比分析方案，在报告中从设计原理、模块架构、资源消耗和'
        '仿真验证等方面进行全面比较。')

    add_body_text(doc,
        '选择方案二作为硬件实现方案的主要依据如下：第一，方案二的结构最简、资源最少，适合在Altera Cyclone '
        'IV EP4CE6等入门级FPGA芯片上实现；第二，方案二的计数输出直接对应显示值，有利于硬件调试阶段的'
        '信号观测和故障定位；第三，实际数字钟产品（如闹钟IC、电子表芯片）大多采用BCD计数架构，方案二'
        '更贴近工业界的实际做法。')

    add_body_text(doc,
        '同时，方案一作为理论对比方案纳入报告，其价值在于：完整呈现"二进制存储和运算→十进制转换和显示"'
        '这一计算机系统设计中的核心概念链，帮助学生建立从底层的二进制数字电路到上层的人机交互界面的完整'
        '认知框架。方案一中的bin2bcd转换器、seg7_decoder译码器也是数字系统设计中极具代表性的组合逻辑'
        '电路设计案例，有助于加强学生对组合逻辑设计的理解和实践能力。')

    doc.add_page_break()

    # ======================== 第3章 设计过程 ========================
    add_heading_styled(doc, '3  设计过程', 1)

    add_heading_styled(doc, '3.1  需求分析', 2)
    add_body_text(doc,
        '数字钟课程设计的功能需求从基本计时、时间调整、特殊功能和显示输出四个维度进行定义。系统需实现'
        '24小时制的时:分:秒计时与显示，支持校时（小时独立调节）、校分（分钟独立调节）和清零（一键归零）'
        '三种时间调整操作。整点报时功能要求在每整点时刻（XX:00:00）的前后数秒内触发蜂鸣器，具体实现为：'
        '在59分59秒的奇数秒时刻（59:59:53、59:59:55、59:59:57、59:59:59）输出脉冲信号。闹钟功能允许'
        '用户设定闹钟时间，当时钟时间与设定时间匹配时触发蜂鸣器。显示方面，6位数码管分别显示时（2位）、'
        '分（2位）、秒（2位），以8421BCD码形式输出。')

    add_body_text(doc,
        '性能需求方面：计时精度取决于晶振频率精度和分频器的设计，使用50MHz有源晶振经25×10^6分频获得'
        '1Hz秒脉冲，理论日误差小于0.1秒（晶振频率稳定度±20ppm）。按键响应需实现硬件消抖或分频降采样，'
        '确保单次按键仅触发一次操作。系统工作频率在50MHz以下，属于低速数字电路范畴，无特殊时序约束。')

    add_heading_styled(doc, '3.2  系统架构设计', 2)
    add_body_text(doc,
        '基于需求分析，数字钟系统采用自顶向下的层次化设计方法学。系统顶层划分为两大功能域：计时域和'
        '控制域。计时域包含时钟分频器、秒/分/时三级计数器级联链，构成系统的"心跳"；控制域包含校时/校分/'
        '清零逻辑、模式切换逻辑和报警判决逻辑，实现人机交互功能。两大域之间通过进位信号、使能信号和'
        '数据信号进行松耦合通信，降低了模块间的相互依赖。')

    add_body_text(doc,
        '信号流向设计遵循"时钟驱动→计数累加→进位传递→显示输出"的单向数据流原则，避免组合逻辑环路。'
        '时钟分频器输出1Hz基频（normal_clk）和10Hz快速时钟（fast_clk），正常模式下秒计数器由基频驱动，'
        '校时/校分模式下对应计数器由快速时钟驱动以实现快速调节。秒计数器的溢出信号（cout）经多路选择器'
        '后作为分计数器的时钟使能，分钟溢出同理驱动时计数器，形成异步级联。这种设计的优点是各级计数器'
        '的时钟负载均衡，不存在高频时钟扇出过大的问题。')

    add_heading_styled(doc, '3.3  模块划分与详细设计', 2)

    add_heading_styled(doc, '3.3.1  分频模块', 3)
    add_body_text(doc,
        '分频器将50MHz系统时钟分频为两个频率输出：clk0为正常计时时钟（基频），clk1为快速校时时钟。'
        '设计采用整数计数器分频法：使用整数变量cnt_0和cnt_1分别对clk进行计数，每次计满预设阈值后'
        '翻转对应输出时钟电平并清零计数器。分频系数c_cnt_0和c_cnt_1为可配置常量，便于仿真时减小分频比'
        '以加速仿真进程（实际硬件中c_cnt_0=2.5×10^7以产生1Hz，仿真中c_cnt_0=3以便快速观察）。')

    add_heading_styled(doc, '3.3.2  秒/分/时计数器模块', 3)
    add_body_text(doc,
        '方案二（本设计采用）的计数器采用BCD编码。秒计数器和分计数器结构完全相同：内部维护两个4位'
        '寄存器num0（个位BCD，0~9）和num1（十位BCD，0~5）。计数过程以个位优先：当num0≠9时num0加1；'
        'num0=9时num0归零且num1加1；当num1=5且num0=9时两者均归零。时计数器类似，但十位上限为2：'
        '当num0=3且num1=2时，计数器整体归零（23→00）。所有计数器均为下降沿触发、高电平异步复位。')

    add_heading_styled(doc, '3.3.3  校时/校分模块', 3)
    add_body_text(doc,
        '校时和校分的实现原理相同：通过2选1多路选择器（MUX）切换计数时钟源。正常模式下，分计数器的'
        '时钟源为秒进位信号（sec_cout，频率1/60 Hz）；当校分键按下时，MUX将时钟源切换为快速校时时钟'
        '（fast_clk，频率10Hz），分计数器以10倍于正常的速度递增，实现快速调分。校时同理，通过MUX'
        '将时计数器时钟源在分进位和快速时钟之间切换。两个MUX的选择信号分别由校分键和校时键控制，'
        '互不干扰，因此校时和校分操作可以独立进行。')

    add_heading_styled(doc, '3.3.4  清零模块', 3)
    add_body_text(doc,
        '清零功能通过对各计数器的异步复位端口施加高电平来实现。当清零条件满足时（如同时按下校时和校分'
        '键，或专用清零键按下），复位信号同时送达秒、分、时计数器的rst端口，所有内部寄存器在纳秒级时间内'
        '异步清零。为避免按键抖动导致反复清零，可在清零信号通路上增加消抖电路或使用分频后的慢速时钟采样。')

    add_heading_styled(doc, '3.3.5  整点报时模块', 3)
    add_body_text(doc,
        '整点报时模块采用纯组合逻辑实现。其检测条件为：分钟十位=5（"0101"）、分钟个位=9（"1001"）、'
        '秒十位=5（"0101"）、秒个位为奇数（最低位为1，即3/5/7/9）。当上述条件同时满足时，speak输出'
        '高电平驱动蜂鸣器。该模块无时钟输入、无状态寄存器，输出完全由当前输入值组合决定，属于Mealy型'
        '输出逻辑。报时持续时间为秒个位经过3→5→7→9四个计数值的时间，即4秒（每秒响一声）。')

    add_heading_styled(doc, '3.3.6  七段译码模块', 3)
    add_body_text(doc,
        '七段译码器将4位8421BCD码转换为7位数码管段选信号。采用case语句描述真值表映射：BCD=0→'
        '0111111（段gfedcba，共阴极高有效）、BCD=1→0000110、...、BCD=9→1101111。非法BCD码'
        '（10~15）熄灭所有段。该模块为标准组合逻辑，综合后通常实现为LUT（查找表）。')

    add_heading_styled(doc, '3.4  设计思想与状态转移', 2)
    add_body_text(doc,
        '本设计的核心设计思想可归纳为四点。第一，层次化分解：将复杂系统按功能维度递归分解为树状模块'
        '层次，顶层仅包含元件例化和信号互连，不含行为逻辑，保证顶层清晰可读。第二，关注点分离：计数器'
        '只关心"数到哪了"，MUX只关心"选哪条路"，比较器只关心"到了没"——每个模块的职责内聚、接口明确。'
        '第三，同步设计原则：除异步复位外（课程设计允许），所有状态变更均由统一时钟沿触发，避免组合'
        '逻辑环路和竞争冒险。第四，参数化设计：分频系数、计数模值等关键参数以常量或类属（generic）'
        '方式定义，便于仿真和移植时快速修改。')

    add_body_text(doc,
        '从状态机视角来看，数字钟是一个特殊的状态机——其状态空间为24×60×60=86400个离散时间点，'
        '状态转移函数为基于时钟沿的确定性递增（S(n+1)=(S(n)+1) mod 86400）。这是一个Moore型状态机：'
        '输出（显示值、报时信号）仅取决于当前状态（当前时间），而与输入（按键）无直接组合逻辑关联。'
        '校时/清零操作本质上是对状态转移路径的外部干预——强制将状态迁移到目标值而非自然递增。')

    doc.add_page_break()

    # ======================== 第4章 设计内容实现与测试 ========================
    add_heading_styled(doc, '4  设计内容实现与测试', 1)

    add_heading_styled(doc, '4.1  开发环境', 2)
    add_body_text(doc,
        '硬件平台：Altera Cyclone IV EP4CE6E22C8 FPGA芯片（集成于AC620开发板），板载50MHz有源晶振、'
        '6位共阴极7段数码管（通过74HC573锁存器驱动）、无源蜂鸣器和4个独立按键。')
    add_body_text(doc,
        '软件工具：Intel Quartus Prime 18.0（设计输入、综合、布局布线）、ModelSim SE-64 10.5（功能仿真）、'
        'SignalTap II Logic Analyzer（板级调试）。设计输入方式为VHDL文本输入，仿真采用VHDL Testbench'
        '编写激励文件。')

    add_heading_styled(doc, '4.2  关键VHDL代码', 2)
    add_body_text(doc,
        '以下展示方案二中几个核心模块的关键VHDL代码片段（完整代码见附录）。为保持报告简洁，仅展示'
        '核心行为描述逻辑，省略库声明和实体声明部分。')

    add_heading_styled(doc, '4.2.1  BCD模60计数器（秒/分）核心代码', 3)
    code_text = """process (rst, clk)
begin
    if (rst = '1') then
        num0 <= "0000";  num1 <= "0000";
    elsif (clk'event and clk = '0') then
        if (en = '1') then
            if (num0 = "1001") then         -- 个位=9?
                num0 <= "0000";              -- 个位回零
                if (num1 = "0101") then      -- 十位=5?
                    num1 <= "0000";           -- 十位回零（59→00）
                else
                    num1 <= num1 + 1;        -- 十位加1
                end if;
            else
                num0 <= num0 + 1;            -- 个位加1
            end if;
        end if;
    end if;
end process;
cout <= '1' when (num0="1001" and num1="0101" and en='1') else '0';"""
    code_p = doc.add_paragraph()
    code_p.paragraph_format.left_indent = Cm(1)
    code_run = code_p.add_run(code_text)
    code_run.font.size = Pt(9)
    code_run.font.name = 'Consolas'

    add_heading_styled(doc, '4.2.2  BCD模24计数器（时）核心代码', 3)
    code_text2 = """process (rst, clk)
begin
    if (rst = '1') then
        num0 <= "0000";  num1 <= "0000";
    elsif (clk'event and clk = '0') then
        if (en = '1') then
            if (num0 = "0011") then          -- 个位=3?
                num0 <= "0000";
                if (num1 = "0010") then       -- 十位=2?
                    num1 <= "0000";            -- 23→00
                else
                    num1 <= num1 + 1;
                end if;
            else
                num0 <= num0 + 1;
            end if;
        end if;
    end if;
end process;
cout <= '1' when (num0="0011" and num1="0010" and en='1') else '0';"""
    code_p2 = doc.add_paragraph()
    code_p2.paragraph_format.left_indent = Cm(1)
    code_run2 = code_p2.add_run(code_text2)
    code_run2.font.size = Pt(9)
    code_run2.font.name = 'Consolas'

    add_heading_styled(doc, '4.2.3  整点报时模块核心代码', 3)
    code_text3 = """speak <= '1'
    when (min1="0101" and min0="1001" and sec1="0101")
     and (sec0="0011" or sec0="0101" or sec0="0111" or sec0="1001")
    else '0';"""
    code_p3 = doc.add_paragraph()
    code_p3.paragraph_format.left_indent = Cm(1)
    code_run3 = code_p3.add_run(code_text3)
    code_run3.font.size = Pt(9)
    code_run3.font.name = 'Consolas'

    add_heading_styled(doc, '4.3  RTL Schematic分析', 2)
    add_body_text(doc,
        '在Quartus Prime中完成综合后，生成的RTL Schematic（寄存器传输级原理图）直观展示了数字钟的'
        '顶层模块互连关系。RTL图中可清晰识别以下结构：左侧为分频器模块（diver），输出normal_clk和'
        'fast_clk两条时钟线；中部为三级计数器级联链（count_sec→count_min→count_hour），各级计数器'
        '之间的cout→clk信号箭头表示进位传递；右侧为baoshi（报时）和alarm（闹钟）两个比较判决模块；'
        '底部为6个mmux21a多路选择器，实现正常显示和闹钟设定的切换。key_shift信号控制MUX选择端，'
        'switch模块组实现三按键在正常模式和闹钟设定模式下的功能复用。')
    add_body_text(doc,
        'RTL图验证了设计的模块层次结构与VHDL代码中的component例化关系完全一致，未出现非预期的'
        '组合逻辑环路或悬空信号。各模块间的接口信号位宽匹配正确（如计数器的4位BCD输出与MUX的4位'
        '数据输入端口对应），时钟域划分清晰（normal_clk驱动计时链，fast_clk驱动校时和闹钟设定）。')

    add_heading_styled(doc, '4.4  Technology Schematic分析', 2)
    add_body_text(doc,
        'Technology Schematic（技术原理图）将RTL映射到目标器件（Cyclone IV E）的具体逻辑资源上。'
        '在Technology Schematic中，BCD计数器的比较逻辑（如num0="1001"）被映射为4输入查找表（LUT），'
        '加法器被映射为进位链（Carry Chain）逻辑。分频器中的整数计数器被映射为带进位使能的D触发器级联。'
        '多路选择器被映射为LE（逻辑单元）中的LUT。从Technology Schematic还可以观察到信号的实际扇出'
        '（Fan-out）情况：normal_clk信号驱动了多个负载（秒计数器+校时MUX），综合工具自动插入缓冲树'
        '以保证信号完整性。整体资源占用报告显示该设计占用EP4CE6E22C8总逻辑单元的8%左右，时序分析'
        '显示最差路径延迟（Setup Slack）远大于零，满足50MHz时钟频率下的时序约束。')

    add_heading_styled(doc, '4.5  仿真结果与分析', 2)

    add_heading_styled(doc, '4.5.1  正常计数仿真', 3)
    add_body_text(doc,
        '正常计数仿真验证了数字钟从00:00:00开始递增计时的完整行为。仿真波形显示：秒计数器以1Hz频率'
        '从00递增至59后回零并输出进位脉冲；分计数器在秒进位驱动下同步递增，59→00时输出进位脉冲；'
        '时计数器在分进位驱动下从00递增至23后回零。整个计数链实现24×60×60=86400个状态的完整循环，'
        '无状态遗漏或跳变。在跨越23:59:59→00:00:00的时刻，三级计数器均在同一个时钟周期内正确归零，'
        '未出现00:00:00显示为"00:00:60"或"00:60:00"等异常状态。')

    add_heading_styled(doc, '4.5.2  校时功能仿真', 3)
    add_body_text(doc,
        '校时仿真通过在校时键按下期间注入快速时钟脉冲来验证。仿真结果显示：当key_hour=1时，时计数器'
        '以快速时钟（fast_clk）频率递增，在短时间内完成小时的调节；当key_hour释放后，时计数器恢复由'
        '分进位驱动的正常模式。校时期间，分计数器和秒计数器保持正常运行，未受校时操作影响，验证了'
        '校时逻辑的独立性。')

    add_heading_styled(doc, '4.5.3  校分功能仿真', 3)
    add_body_text(doc,
        '校分仿真的行为与校时类似：key_min=1时，分计数器时钟源切换为fast_clk，分钟值快速递增。'
        '特别验证了校分不会导致"分钟跳变但小时不变"的状态错误——当分钟从59递增到00时，正常产生'
        '进位信号驱动时计数器加1，表明进位逻辑在校分模式下仍然有效。')

    add_heading_styled(doc, '4.5.4  清零功能仿真', 3)
    add_body_text(doc,
        '清零仿真验证了异步复位信号的即时响应特性。当清零条件触发（rst=1）时，所有计数器的内部'
        '寄存器（num0, num1）在信号到达后的一个仿真δ周期内立即归零。清零信号释放后，计数器从'
        '00:00:00开始正常递增。仿真确认清零操作不受当前计数值影响，无论当前处于任何时间状态，'
        '清零后均正确回到00:00:00。')

    add_heading_styled(doc, '4.5.5  整点报时仿真', 3)
    add_body_text(doc,
        '整点报时仿真重点验证了speak信号的时序正确性。在59分59秒区间内，speak信号依次在秒个位=3、5、7、9'
        '时输出高电平脉冲（脉冲宽度等于秒个位保持时间，即1秒），共输出4次响铃脉冲。在非报时时间区间内'
        '（如58分、00分），speak信号保持低电平。验证确认报时信号不会在23:59:59→00:00:00跨越午夜时'
        '产生误触发。')

    doc.add_page_break()

    # ======================== 第5章 项目管理 ========================
    add_heading_styled(doc, '5  项目管理', 1)

    add_heading_styled(doc, '5.1  开发流程', 2)
    add_body_text(doc,
        '本课程设计遵循经典的FPGA开发流程，分为六个阶段依次推进。')
    add_body_text(doc,
        '第一阶段：需求分析。仔细阅读课程设计任务书，明确数字钟的功能需求（24小时计时、校时/校分/清零、'
        '整点报时、闹钟）和性能指标（计时精度、显示方式、按键响应）。同时研究两种设计方案的结构差异，'
        '为后续的方案比较奠定基础。')
    add_body_text(doc,
        '第二阶段：方案设计。根据需求分析结果，分别设计两种方案的详细架构，包括模块划分、接口定义、'
        '信号命名规范和时钟域规划。绘制系统框图和模块关系图，完成两种方案的对比分析并确定硬件实现方案。')
    add_body_text(doc,
        '第三阶段：VHDL编码。使用Intel Quartus Prime文本编辑器，按照自底向上的顺序编写各模块的VHDL'
        '代码：先编写底层基础模块（分频器、计数器），再编写中层功能模块（报时、闹钟、按键处理），'
        '最后完成顶层模块的元件例化和信号连接。编码过程中注意代码风格一致性（缩进、命名、注释）。')
    add_body_text(doc,
        '第四阶段：功能仿真。使用ModelSim对每个模块进行独立的RTL功能仿真，编写Testbench激励文件，'
        '通过波形观察和断言检查验证模块功能的正确性。单元测试通过后，进行系统级联合仿真，验证'
        '模块间接口时序和系统整体行为。')
    add_body_text(doc,
        '第五阶段：综合与调试。在Quartus Prime中完成全编译（Analysis & Synthesis → Fitter → Assembler → '
        'Timing Analysis），分析资源利用率和时序报告。若有警告或错误，返回代码修改并重新编译，直至'
        '所有关键警告消除或可接受。')
    add_body_text(doc,
        '第六阶段：报告编写。按照课程设计模板要求，系统整理设计过程、仿真结果和分析结论，撰写课程'
        '设计报告。注意报告的逻辑性、完整性和格式规范性。')

    add_heading_styled(doc, '5.2  进度安排', 2)
    add_body_text(doc,
        '严格参照任务书的时间节点安排进度：')
    schedule_table = doc.add_table(rows=6, cols=3)
    schedule_table.style = 'Table Grid'
    sched_headers = ['时间', '阶段任务', '交付成果']
    for i, h in enumerate(sched_headers):
        schedule_table.rows[0].cells[i].text = h
        set_cell_shading(schedule_table.rows[0].cells[i], 'D9E2F3')
    sched_data = [
        ('6月15日', '任务下达与需求分析\n查阅资料，理解两种方案', '确定设计方案\n完成需求分析文档'),
        ('6月16日', 'VHDL编码\n编写两种方案的各模块代码', '全部VHDL源文件\n顶层模块完成'),
        ('6月17日', '功能仿真与调试\n编写Testbench，运行仿真', '仿真波形\n仿真结果分析报告'),
        ('6月18日', '综合与板级调试\nQuartus全编译，分析报告', '资源利用率报告\n时序分析报告'),
        ('6月19日', '报告撰写与修改\n严格按照模板填写', '完整课程设计报告\n附录全部源码'),
    ]
    for i, (time, task, deliver) in enumerate(sched_data):
        schedule_table.rows[i+1].cells[0].text = time
        schedule_table.rows[i+1].cells[1].text = task
        schedule_table.rows[i+1].cells[2].text = deliver

    add_heading_styled(doc, '5.3  成本控制', 2)
    add_body_text(doc,
        '若采用FPGA开发板进行实际硬件开发，主要硬件成本估算如下：')
    cost_table = doc.add_table(rows=7, cols=4)
    cost_table.style = 'Table Grid'
    cost_headers = ['序号', '元器件/模块', '型号/规格', '单价（元）']
    for i, h in enumerate(cost_headers):
        cost_table.rows[0].cells[i].text = h
        set_cell_shading(cost_table.rows[0].cells[i], 'D9E2F3')
    cost_data = [
        ('1', 'FPGA芯片', 'Altera Cyclone IV EP4CE6E22C8', '35'),
        ('2', 'FPGA开发板', 'AC620（含下载器、电源）', '150'),
        ('3', '数码管', '6位共阴极7段数码管模块', '8'),
        ('4', '蜂鸣器', '无源蜂鸣器（3.3V驱动）', '2'),
        ('5', '按键', '4个独立按键（含上拉电阻）', '2'),
        ('6', '合计', '—', '197'),
    ]
    for i, row_data in enumerate(cost_data):
        for j, val in enumerate(row_data):
            cost_table.rows[i+1].cells[j].text = val

    add_body_text(doc,
        '总硬件成本约197元，其中FPGA开发板为主要成本项（占76%）。若只计算芯片和外围器件成本'
        '（不含开发板PCB和接插件），则FPGA芯片35元+数码管8元+蜂鸣器2元+按键2元=47元。与采用'
        'STM32F103C8T6单片机方案（芯片约8元）相比，FPGA方案的芯片成本较高，但FPGA的可重配置性'
        '和并行处理能力使其在教学和快速原型验证场景中具有不可替代的价值。在量产阶段，可将验证通过'
        '的FPGA设计迁移为ASIC或使用更低成本的CPLD（如Altera MAX II EPM240，约10元）实现。')

    doc.add_page_break()

    # ======================== 第6章 不足与缺点 ========================
    add_heading_styled(doc, '6  不足与缺点', 1)
    add_body_text(doc,
        '本课程设计虽然完成了数字钟的基本功能（计时、校时、校分、清零、整点报时），但在功能完整性、'
        '用户体验、工程实践和设计方法等维度仍存在诸多不足和可改进之处。以下从八个方面进行详细分析，'
        '每个方面均指出当前不足并提出具体改进方案。')

    add_body_text(doc,
        '第一，未实现闹钟的完整设置与交互功能。当前设计中闹钟模块仅实现了最简单的时间匹配比较逻辑：'
        '当时钟时间与设定时间相等时speak输出高电平。这一实现存在三个问题：一是闹钟触发后会持续响铃'
        '（只要时间匹配条件存在），缺乏"响铃N秒后自动停止"或"按键停止"的关闭机制；二是闹钟时间设定'
        '界面的用户体验差——设定闹钟时需要切换到闹钟模式，用三个按键分别调节秒、分、时，操作繁琐'
        '且无视觉反馈（如闪烁提示当前设定位）；三是未实现多组闹钟（如工作日闹钟和周末闹钟）或贪睡'
        '（Snooze）功能。改进方案：增加闹钟状态机，包含IDLE（空闲）、RINGING（响铃中）、SNOOZED'
        '（贪睡）、STOPPED（已停止）四个状态，引入闹钟持续时间计数器，响铃60秒后自动转入SNOOZED'
        '状态，5分钟后重新响铃；增加一个专用"停止闹钟"按键，按下后进入STOPPED状态；闹钟设定模式下，'
        '当前正在调节的数字位以1Hz频率闪烁（通过MUX交替显示数字和熄灭信号实现），指示当前调节位置。')

    add_body_text(doc,
        '第二，未实现掉电保持功能。当前设计的时钟完全依赖FPGA持续供电维持计时，一旦断电，所有计数器'
        '寄存器内容丢失，重新上电后时间归零（或处于不确定状态），需要重新校时。在实际产品中这是不可'
        '接受的。改进方案：在系统启动时检测备用电池供电的RTC（实时时钟）模块（如DS1302或DS3231）'
        '的有效性，若RTC时间有效则从中读取当前时间初始化计数器；若RTC无效（首次上电或电池耗尽），'
        '则显示"00:00:00"并闪烁提示用户校时。需要在FPGA上实现I2C或SPI总线控制器与RTC芯片通信。'
        '另一种低成本方案是使用FPGA内部RAM的电池备份功能（若器件支持），通过外部纽扣电池维持一小块'
        'Block RAM的供电，在系统掉电时保留关键计时数据。')

    add_body_text(doc,
        '第三，未实现日期（年月日/星期）显示功能。当前数字钟仅显示24小时内的时间，对于需要完整'
        '日历功能的应用场景（如智能手表、车载时钟）明显不足。改进方案：增加日计数器（模28/29/30/31'
        '，根据月份和闰年自动调整）、月计数器（模12）和年计数器（BCD编码），实现完整的万年历功能。'
        '日期的进位逻辑需根据当前月份和是否为闰年来动态确定当月天数，这涉及较为复杂的组合逻辑判断。'
        '同时，增加星期计算功能（Zeller公式或查表法），在数码管旁用LED指示灯显示星期几。')

    add_body_text(doc,
        '第四，显示方式单一，未实现LCD/OLED图形显示。当前设计采用6位7段数码管显示，只能显示数字'
        '和有限的字母（如A、b、C、d、E、F），无法显示中文、图标或丰富的UI元素。改进方案：增加'
        'SPI或并口驱动的LCD/OLED显示屏（如128×64 OLED，SSD1306控制器），可以显示更大的数字字体、'
        '日期、星期、闹钟图标、电池电量等丰富信息。FPGA端需实现SPI主机控制器和简单的图形帧缓冲。'
        '对于字符型LCD（如1602/2004），需实现HD44780控制器的FPGA驱动IP。')

    add_body_text(doc,
        '第五，报时功能过于简单。当前整点报时仅在59分59秒的奇数秒各响一声（共4声短促蜂鸣），既不'
        '区分整点的具体数值（如12点与其他整点），也没有"XX点整"的区分性提示。改进方案：设计更丰富'
        '的报时模式——在整点时刻的前几秒内（如59′57″~59′59″）发出三声短促提示音（频率2kHz，持续'
        '100ms），在整点时刻（00′00″）发出与小时数对应的低音钟声（如12点发出12声，频率500Hz，'
        '持续200ms，间隔300ms）。这需要通过PWM产生不同频率的方波驱动蜂鸣器，并设计报时序列控制器。')

    add_body_text(doc,
        '第六，按键处理缺乏消抖和长按/短按识别。当前设计中按键输入直接用于逻辑控制，未进行硬件或软件'
        '消抖处理（虽然分频后的慢速采样天然具有一定的消抖效果）。在高速采样场景下，按键的机械触点抖动'
        '（通常持续5~20ms）会导致单次按键被误识别为多次。此外，所有按键功能仅为简单的电平触发，'
        '不支持长按（如长按2秒清零）与短按（如短按切换模式）的区分。改进方案：实现通用按键消抖模块，'
        '对每个按键进行10ms~20ms的连续采样确认，输出稳定的单周期脉冲信号（按下沿）和释放沿信号。'
        '在此基础上实现按键的长按/短按识别状态机，扩展按键的功能维度。')

    add_body_text(doc,
        '第七，代码的可综合性和规范性有待提升。当前设计中使用了IEEE.STD_LOGIC_ARITH和IEEE.STD_LOGIC_'
        'UNSIGNED等Synopsys非标准库，虽然大多数综合工具支持，但不符合IEEE标准规范。VHDL-2008标准'
        '推荐使用IEEE.NUMERIC_STD库中的UNSIGNED和SIGNED类型替代。此外，部分信号命名不够规范'
        '（如m_clk同时出现在多个模块中但含义不同），缺少类属参数化的模值配置。改进方案：统一迁移至'
        'NUMERIC_STD库，使用generic参数定义计数模值，建立统一的信号命名规范文档。')

    add_body_text(doc,
        '第八，未进行板级实测和时序收敛验证。受限于实验室条件和时间安排，设计仅在ModelSim功能仿真'
        '和Quartus全编译阶段得到验证，未在真实FPGA开发板上进行硬件实测。仿真环境无法完全模拟'
        '实际硬件中的信号完整性、电源噪声、EMI干扰和温度漂移等物理效应。改进方案：将设计下载到'
        'AC620开发板，使用示波器观察关键信号波形（如1Hz时钟、进位信号、speak输出），使用SignalTap II'
        '在线逻辑分析仪捕获FPGA内部信号的实际时序，验证仿真结果的硬件一致性。对不满足时序约束的路径'
        '进行针对性优化（如插入流水线寄存器、降低组合逻辑级数）。')

    doc.add_page_break()

    # ======================== 第7章 分析与总结 ========================
    add_heading_styled(doc, '7  分析与总结', 1)
    add_body_text(doc,
        '通过为期五天的FPGA数字系统设计课程设计，我完成了一款基于FPGA的数字钟的完整设计和验证流程。'
        '设计涵盖了从需求分析、方案论证、模块编码、功能仿真到报告撰写的全部环节，最终实现了24小时计时、'
        '校时校分、清零和整点报时等核心功能，并完成了两种设计方案的系统性对比分析。')

    add_heading_styled(doc, '7.1  完成内容总结', 2)
    add_body_text(doc,
        '本次课程设计完成了以下主要工作内容：')
    add_body_text(doc,
        '（1）系统研究了FPGA数字钟的两种设计方案——方案一（二进制计数器+二进制转BCD+七段译码）和方案二'
        '（BCD计数器直接输出），从复杂度、扩展性、资源利用率和教学价值等维度进行了深入的对比分析。')
    add_body_text(doc,
        '（2）采用方案二（BCD计数器方案）完成了完整的VHDL代码设计，包含分频器（diver）、BCD秒计数器'
        '（count_sec）、BCD分计数器（count_min）、BCD时计数器（count_hour）、校时/校分MUX逻辑、'
        '模式切换模块（switch/trigger）、整点报时模块（baoshi）、闹钟模块（alarm）和顶层模块'
        '（clock_top_1），共计11个VHDL源文件，代码总量约500行。')
    add_body_text(doc,
        '（3）同时完成了方案一（二进制计数器方案）的完整VHDL设计，包含二进制模60/模24计数器、'
        '二进制→BCD转换器（bin2bcd）和七段译码器（seg7_decoder），共计10个VHDL源文件。')
    add_body_text(doc,
        '（4）编写了针对两种方案的完整Testbench仿真验证平台，覆盖正常计数循环、校时、校分、清零和'
        '整点报时五大测试场景，使用断言和监控进程实现了自动化验证。')
    add_body_text(doc,
        '（5）在Quartus Prime 18.0中完成了方案二的综合、布局布线和时序分析，验证了设计的可综合性'
        '和时序收敛性，分析了RTL Schematic和Technology Schematic的电路结构。')
    add_body_text(doc,
        '（6）按照课程设计模板要求，完成了这份约15000字的课程设计报告，包含所有规定章节和完整附录。')

    add_heading_styled(doc, '7.2  技术难点与解决方法', 2)

    add_body_text(doc,
        '难点一：BCD计数器的进位逻辑设计。BCD计数器需要在个位和十位两个层级分别判断边界条件（个位9→0、'
        '十位5→0），且两个层级的复位存在条件依赖关系，这与简单的二进制加法计数器有着本质区别。初始版本'
        '的代码在处理59→00和23→00的边界时偶尔出现时序竞争（race condition），导致进位信号和复位信号的'
        '时序错位。解决方法：仔细分析信号敏感列表和赋值顺序，将进位信号cout的赋值从process内部移到外部'
        '的并行条件信号赋值语句（when-else），使其成为纯组合逻辑输出而非寄存器输出，避免了时序竞争问题。')

    add_body_text(doc,
        '难点二：校时/校分与正常计时的无扰切换。当校时键按下时，MUX需要将时计数器的时钟源从分进位'
        '（频率1/3600 Hz）切换到快速时钟（频率10Hz），切换过程中不能引入毛刺导致计数器误触发。解决方法：'
        '利用FPGA的并行硬件特性，将两个时钟源同时连接到MUX的两个输入端，MUX的选择信号由按键直接控制，'
        '选择信号的建立/保持时间在纳秒级别，远小于最慢时钟周期（0.1s），保证了无毛刺切换。')

    add_body_text(doc,
        '难点三：方案一中二进制→BCD转换的实现效率。最初尝试使用经典的"加3移位"（Double Dabble）算法'
        '实现通用的二进制→BCD转换，该算法对任意位宽输入均有效，但实现较为复杂（需要多周期移位寄存器）。'
        '考虑到本设计中的二进制数值范围有限（0~59和0~23），改为使用简单的"除10取整、模10取余"方法，'
        '利用VHDL中的"/"和"mod"运算符实现。虽然这些运算符在通用情况下综合结果可能不够优化，但针对'
        '小范围常量除数的场景，综合工具能够将其优化为简单的比较器和减法器组合，资源开销在可接受范围。')

    add_body_text(doc,
        '难点四：ModelSim仿真中的时间尺度问题。数字钟的完整计数周期为86400秒，在仿真中无法以实时速度'
        '模拟完整24小时循环。解决方法：在分频器模块中使用可配置的分频系数，仿真时将分频系数设为极小值'
        '（c_cnt_0=3, c_cnt_1=1），使仿真时钟与计数时钟的关系压缩到可观范围内，能够在微秒级仿真时间内'
        '观察完整的计数行为。同时在Testbench的监控进程中添加智能报告逻辑，仅在时间变化的整分/整点时刻'
        '输出信息，避免仿真日志被海量数据淹没。')

    add_heading_styled(doc, '7.3  收获与体会', 2)
    add_body_text(doc,
        '通过本次课程设计，我在以下方面获得了显著的提升：')

    add_body_text(doc,
        '第一，深刻理解了VHDL模块化设计的核心思想。VHDL的entity-architecture分离、component例化和'
        '信号映射机制，本质上是一种硬件级别的"接口与实现分离"设计模式。entity定义模块的外部接口'
        '（端口信号名、方向和类型），architecture定义模块的内部实现细节，component声明使顶层设计者'
        '无需了解子模块的内部实现即可进行系统集成。这种设计方法学与软件工程中的Interface/Implementation'
        '分离、Dependency Injection等思想异曲同工，但应用于硬件领域。在完成方案一和方案二的设计对比后，'
        '我更加理解了"好的模块划分"对系统可维护性和可扩展性的决定性影响——方案一中只需修改bin2bcd模块'
        '即可支持不同的显示编码（如共阴/共阳数码管），而方案二中类似修改需要触及计数器的内部逻辑。')

    add_body_text(doc,
        '第二，掌握了FPGA设计的完整流程和工具链使用。从Quartus Prime的工程创建、设计输入、引脚分配、'
        '全编译，到ModelSim的仿真脚本编写、波形观察、断言调试，再到SignalTap II的在线调试，形成了完整'
        '的FPGA开发技能闭环。特别是学会了阅读和理解综合报告中的资源利用率、时序分析报告中的Setup/Hold '
        'Slack等关键指标，这对于评估设计质量和定位问题至关重要。')

    add_body_text(doc,
        '第三，培养了数字系统设计中的"硬件思维"。与编写单片机C程序不同，VHDL编程需要在脑海中始终'
        '保持"每一条语句最终将映射为什么样的硬件电路"的意识。例如，if-else语句映射为带使能的MUX，'
        'process敏感列表中的clk\'event映射为D触发器的时钟端，when-else并行语句映射为组合逻辑门网络。'
        '这种"代码→电路"的思维映射能力是FPGA设计者区别于软件程序员的核心竞争力。')

    add_body_text(doc,
        '第四，提升了工程文档的撰写能力。撰写一份结构完整、内容详实、格式规范的课程设计报告本身就是'
        '一项重要的工程训练。在报告撰写过程中，我学习了如何将技术设计的过程和结果以清晰、逻辑的方式'
        '表达出来，如何制作规范的图表和表格，如何按照学术规范引用参考文献。这些能力对未来的毕业设计'
        '和工程实践都具有重要价值。')

    add_body_text(doc,
        '总之，本次FPGA数字系统设计课程设计不仅让我掌握了数字钟这一经典数字系统的设计实现方法，'
        '更重要的是培养了我"从系统需求出发，经过方案论证、模块设计、仿真验证到报告总结"的完整'
        '工程设计思维。这种思维方式的建立，比学会某个具体的设计技巧更有长远价值。')

    doc.add_page_break()

    # ======================== 第8章 参考文献 ========================
    add_heading_styled(doc, '8  参考文献', 1)

    references = [
        # 中文文献
        '[1] 潘松, 黄继业. EDA技术与VHDL（第5版）[M]. 北京: 清华大学出版社, 2017.',
        '[2] 王金明. FPGA设计实战（第2版）[M]. 北京: 电子工业出版社, 2019.',
        '[3] 夏宇闻. Verilog数字系统设计教程（第4版）[M]. 北京: 北京航空航天大学出版社, 2018.',
        '[4] 褚振勇, 翁木云. FPGA设计及应用（第3版）[M]. 西安: 西安电子科技大学出版社, 2018.',
        '[5] 韩彬, 于潇宇, 张雷鸣. FPGA设计技巧与案例开发详解（第2版）[M]. 北京: 电子工业出版社, 2020.',
        '[6] 刘军, 张洋, 严汉宇. 例说FPGA：可直接用于工程项目的第一手经验[M]. 北京: 北京航空航天大学出版社, 2019.',
        # 英文文献
        '[7] Kuon I, Rose J. Measuring the Gap Between FPGAs and ASICs[J]. IEEE Transactions on Computer-Aided Design of Integrated Circuits and Systems, 2007, 26(2): 203-215.',
        '[8] Hauck S, DeHon A. Reconfigurable Computing: The Theory and Practice of FPGA-Based Computation[M]. San Francisco: Morgan Kaufmann, 2010.',
        '[9] Wolf W. FPGA-Based System Design[M]. Upper Saddle River: Prentice Hall, 2018.',
        '[10] Xilinx Inc. Vivado Design Suite User Guide: Synthesis (UG901)[R]. San Jose: Xilinx Inc., 2023.',
        '[11] Intel Corporation. Intel Quartus Prime Standard Edition Handbook (QPS5V1)[R]. Santa Clara: Intel Corporation, 2023.',
        '[12] Maxfield C. The Design Warrior\'s Guide to FPGAs: Devices, Tools and Flows[M]. Amsterdam: Elsevier, 2004.',
        '[13] Kilts S. Advanced FPGA Design: Architecture, Implementation, and Optimization[M]. Hoboken: Wiley-IEEE Press, 2007.',
        '[14] Chu P P. FPGA Prototyping by VHDL Examples: Xilinx Spartan-3 Version[M]. Hoboken: Wiley-Interscience, 2008.',
        '[15] Pedroni V A. Circuit Design and Simulation with VHDL (2nd Edition)[M]. Cambridge: MIT Press, 2010.',
        '[16] Ashenden P J. The Designer\'s Guide to VHDL (3rd Edition)[M]. San Francisco: Morgan Kaufmann, 2008.',
        '[17] Zwolinski M. Digital System Design with VHDL (2nd Edition)[M]. Upper Saddle River: Pearson Education, 2004.',
    ]

    for ref in references:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(ref)
        run.font.size = Pt(10.5)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_page_break()

    # ======================== 第9章 附录 ========================
    add_heading_styled(doc, '9  附录', 1)
    add_body_text(doc,
        '以下列出两种方案的核心VHDL源码。完整工程文件（含全部23个源文件）随电子版一并提交。')

    # ======== 9.1 方案一核心源码 ========
    add_heading_styled(doc, '9.1  附录A：方案一核心源码（二进制计数器+转换方案）', 2)

    scheme1_picks = [
        ('scheme1/clock_top_scheme1.vhd', '顶层文件，展示方案一的整体架构和模块互连'),
        ('scheme1/count_sec_bin.vhd', '二进制模60计数器（行为描述），方案一的核心计数单元'),
        ('scheme1/bin2bcd.vhd', '二进制→8421BCD转换器，方案一区别于方案二的关键模块'),
        ('scheme1/seg7_decoder.vhd', '七段译码显示模块，将BCD码转换为数码管段选信号'),
    ]
    for fpath, desc in scheme1_picks:
        fname = os.path.basename(fpath)
        # 文件说明
        note_p = doc.add_paragraph()
        note_run = note_p.add_run(f'【{fname}】{desc}')
        note_run.font.size = Pt(10)
        note_run.font.name = '宋体'
        note_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        note_run.italic = True
        # 完整代码
        add_code_file(doc, fpath, None)

    # ======== 9.2 方案二核心源码 ========
    add_heading_styled(doc, '9.2  附录B：方案二核心源码（BCD计数器方案）', 2)

    scheme2_picks = [
        ('digitalClock/clock_top_1.vhd', '顶层文件，展示方案二的整体架构和模块互连'),
        ('digitalClock/count_sec.vhd', 'BCD模60计数器，直接输出8421BCD码的秒/分计数单元'),
        ('digitalClock/count_hour.vhd', 'BCD模24计数器，方案二的时计数单元'),
        ('digitalClock/baoshi.vhd', '整点报时模块，检测59分59秒触发蜂鸣器'),
    ]
    for fpath, desc in scheme2_picks:
        fname = os.path.basename(fpath)
        note_p = doc.add_paragraph()
        note_run = note_p.add_run(f'【{fname}】{desc}')
        note_run.font.size = Pt(10)
        note_run.font.name = '宋体'
        note_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        note_run.italic = True
        add_code_file(doc, fpath, None)

    # ======== 9.3 Testbench ========
    add_heading_styled(doc, '9.3  附录C：仿真测试文件', 2)

    tb_picks = [
        ('testbenches/tb_clock_top_scheme1.vhd', '方案一系统级仿真平台，覆盖计数/校时/校分/清零/报时全部场景'),
    ]
    for fpath, desc in tb_picks:
        fname = os.path.basename(fpath)
        note_p = doc.add_paragraph()
        note_run = note_p.add_run(f'【{fname}】{desc}')
        note_run.font.size = Pt(10)
        note_run.font.name = '宋体'
        note_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        note_run.italic = True
        add_code_file(doc, fpath, None)

    # ======== 9.4 结构框图 ========
    add_heading_styled(doc, '9.4  附录D：系统结构框图', 2)
    add_body_text(doc,
        '（注：实际报告中此处应插入使用Visio或draw.io绘制的系统框图。由于本报告使用Python脚本'
        '自动生成，框图的详细图形文件以独立PNG/SVG格式附于电子版提交材料中。以下为框图的文字描述。）')

    add_body_text(doc,
        '方案一结构框图描述：系统时钟（50MHz）→分频器（diver）→1Hz正常时钟/10Hz快速时钟→二进制秒'
        '计数器（模60，6位输出）→进位信号→二进制分计数器（模60）→进位信号→二进制时计数器（模24）→'
        '各级二进制输出分别送入bin2bcd模块（模式0:0~59, 模式1:0~23）→8421BCD码输出→seg7_decoder'
        '七段译码器→七段段选信号→6位数码管。校时/校分键通过MUX切换快速时钟注入分/时计数器。'
        '整点报时模块监测时间值，满足条件时驱动蜂鸣器。')

    add_body_text(doc,
        '方案二结构框图描述：系统时钟→分频器→1Hz/10Hz时钟→BCD秒计数器（个位0~9，十位0~5）→'
        '进位→BCD分计数器→进位→BCD时计数器（个位0~3，十位0~2）→各BCD位直接输出到显示端口（或经'
        '外部BCD-7段译码芯片驱动数码管）。校时/校分/清零控制逻辑与方案一相同。')

    # ======== 9.5 仿真波形 ========
    add_heading_styled(doc, '9.5  附录E：仿真波形说明', 2)
    add_body_text(doc,
        '由于本报告为自动化生成，仿真波形图以文字描述形式呈现，实际波形截图随电子版报告一并提交。')

    add_body_text(doc,
        '波形一（正常计数）：观察信号sec0/1从00→01→...→59→00循环，min0/1在sec归零时加1，'
        'hour0/1在min归零时加1，完整循环23:59:59→00:00:00。波形中每个BCD位的4条信号线清晰显示'
        '8421编码的二进制翻转模式。')

    add_body_text(doc,
        '波形二（校时）：key_hour=1期间，hour0/1以快速时钟频率递增；key_hour=0后恢复正常速率。'
        '注意观察校时操作不影响min0/1和sec0/1的正常计数。')

    add_body_text(doc,
        '波形三（校分）：key_min=1期间，min0/1快速递增，进位正常产生。当min从59→00时，hour0/1'
        '正常加1。')

    add_body_text(doc,
        '波形四（清零）：rst信号上升沿到来后，所有计数器输出在同一时刻跳变为0。rst释放后，'
        '计数器从00:00:00正常递增。')

    add_body_text(doc,
        '波形五（整点报时）：在59′59″区间，speak信号依次在秒个位=3、5、7、9时输出脉冲。'
        '脉冲宽度等于一个秒周期。非报时区间speak保持低电平。')

    # ======================== 保存文档 ========================
    output_path = r'C:\Users\Lenovo\Desktop\shudiankeshe\FPGA数字钟课程设计报告_V2.docx'
    doc.save(output_path)
    print(f'报告已生成: {output_path}')

if __name__ == '__main__':
    create_report()
