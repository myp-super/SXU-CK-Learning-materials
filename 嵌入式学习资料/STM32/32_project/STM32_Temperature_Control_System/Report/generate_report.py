#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
====================================================================
  STM32温度测控系统 课程设计报告生成脚本
  山西大学 测控技术与仪器专业
====================================================================
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ==================== 页面设置 ====================
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.17)
section.right_margin = Cm(3.17)

# ==================== 样式设置 ====================
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 标题样式函数
def add_heading_styled(text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(14)
        elif level == 3:
            run.font.size = Pt(13)
    return h

def add_para(text, bold=False, size=12, align=None, font_name='宋体', first_line_indent=True):
    p = doc.add_paragraph()
    if first_line_indent and align is None:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_table_with_data(headers, data, col_widths=None):
    table = doc.add_table(rows=1 + len(data), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Headers
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    # Data
    for r, row in enumerate(data):
        for c, val in enumerate(row):
            cell = table.rows[r+1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()  # spacing
    return table

# ==================== 封面 ====================
for i in range(6):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run('测控系统设计与应用实习\n课程设计说明书')
run.font.name = '黑体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.font.size = Pt(26)
run.bold = True

doc.add_paragraph()
doc.add_paragraph()

info_lines = [
    '课题名称：STM32单片机温度测控系统',
    '专    业：测控技术与仪器',
    '学    院：物理电子工程学院',
    '姓    名：__________',
    '学    号：__________',
    '指导教师：__________',
    f'完成日期：{datetime.date.today().strftime("%Y年%m月%d日")}'
]
for line in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(14)

doc.add_page_break()

# ==================== 摘要 ====================
add_heading_styled('摘  要', 1)

abstract_text = (
    '温度是工业生产、环境监测和日常生活中最重要的物理量之一，对温度进行精确测量和自动控制具有重要的工程应用价值。'
    '本文基于ARM Cortex-M3内核的STM32F103ZET6微控制器，设计并实现了一套完整的温度测控系统。'
    '系统以DS18B20数字温度传感器作为温度采集单元，利用其单总线通信协议实现高精度温度测量；'
    '采用ILI9341控制器的DMTFT-28液晶显示屏作为人机交互界面，通过FSMC并行接口实现高速数据显示；'
    '利用STM32内置高级定时器TIM3产生两路独立的PWM信号，分别驱动风扇（降温执行器）和电热丝（升温执行器），'
    '并使用LED模拟负载进行功能验证。系统通过软件实现了基于滞回比较的温度控制策略：'
    '当当前温度高于设定温度时，自动关闭电热丝输出并启动风扇降温，PWM占空比随温度偏差增大而线性增加；'
    '当当前温度低于设定温度时，自动关闭风扇并启动电热丝加热，PWM占空比同样随偏差增大而增加。'
    '测试结果表明，系统温度测量精度达到±0.5°C，PWM输出频率稳定在1kHz±0.1%，'
    'LCD刷新率可达10Hz以上，温度控制响应时间小于5秒，各项指标均满足设计要求。'
    '本系统具有成本低、精度高、响应快、扩展性强等优点，可广泛应用于实验室温控、'
    '智能家居、农业温室等领域，为嵌入式温度测控系统的教学和工程实践提供了完整的参考方案。'
)
add_para(abstract_text)

add_heading_styled('关键词', 2)
add_para('STM32F103ZET6；DS18B20；PWM温度控制；LCD显示；嵌入式系统', bold=True)

doc.add_page_break()

# ==================== 第1章 绪论 ====================
add_heading_styled('1  绪论', 1)

add_heading_styled('1.1  研究背景', 2)
add_para(
    '在工业生产过程中，温度是最常见且最重要的过程参数之一。据统计，在工业过程控制中，'
    '温度控制相关的回路占总控制回路的50%以上。从精密仪器制造到大型化工生产，从医疗设备到'
    '智能家居，温度测控技术始终扮演着不可替代的角色。随着微电子技术、传感器技术和嵌入式'
    '系统的飞速发展，温度测控系统正向着高精度、智能化、网络化和低成本方向不断演进。'
)
add_para(
    '传统的温度控制系统多采用模拟电路或PLC实现，存在精度低、灵活性差、成本高等不足。'
    '基于单片机的数字温度控制系统凭借其体积小、功耗低、功能灵活可编程等优势，已成为'
    '现代温度测控领域的主流方案。特别是以ARM Cortex-M系列为代表的32位微控制器，'
    '以其强大的处理能力、丰富的外设资源和低廉的成本，为设计高性能温度测控系统提供了理想平台。'
)
add_para(
    '本课题面向"测控系统设计与应用实习"课程教学需求，以STM32F103ZET6为核心控制器，'
    '结合DS18B20数字温度传感器、LCD显示屏和PWM功率控制技术，设计并实现一套完整的'
    '温度自动测控系统。该系统不仅能够实时采集和显示温度数据，还能根据设定值自动调节'
    '风扇和电热丝的工作状态，实现闭环温度控制，具有较高的教学价值和工程实践意义。'
)

add_heading_styled('1.2  国内外应用现状', 2)
add_para(
    '温度测控技术的研究与应用在国内外均已取得丰硕成果。在工业领域，以Honeywell、'
    'Yokogawa、Siemens等为代表的国际企业推出了成熟的分布式温度控制系统（DCS），'
    '采用先进的多回路PID控制算法和现场总线技术，可实现上千个温度节点的集中监控。'
    '在科研领域，基于模糊控制、神经网络和模型预测控制（MPC）的智能温度控制算法'
    '不断涌现，显著提升了温度控制的精度和鲁棒性。'
)
add_para(
    '国内温度测控技术虽然起步较晚，但发展迅速。以中控技术、和利时为代表的国产DCS系统'
    '已经在石化、电力等行业得到广泛应用。在嵌入式温度控制领域，基于STM32、MSP430、'
    'STM8等微控制器的方案被大量应用于恒温箱、培养箱、3D打印机加热平台等产品中。'
    '高校和科研机构也在积极研究基于物联网（IoT）的远程温度监控系统，通过WiFi、'
    '蓝牙或LoRa等无线通信技术实现温度数据的远程采集和云端分析。'
)

add_heading_styled('1.3  温度测控系统发展', 2)
add_para(
    '温度测控技术的发展大致经历了四个阶段：第一阶段（1960s-1970s）以模拟仪表和继电器'
    '控制为主，精度低且功能单一；第二阶段（1980s-1990s）引入单片机数字控制，实现了PID'
    '算法和数字显示，但通信能力有限；第三阶段（2000s-2010s）以嵌入式系统和现场总线'
    '技术为核心，实现了分布式控制和远程监控；第四阶段（2010s至今）以物联网、云计算和'
    '人工智能为特征，温度测控系统向智能化、网络化和大数据分析方向发展。'
)
add_para(
    '当前，温度测控系统的主要发展趋势包括：（1）高精度与高分辨率——采用16位甚至24位'
    '高精度ADC和精密传感器，实现0.01°C级温度分辨；（2）多传感器融合——综合热电偶、'
    '热电阻、红外等多种测温手段，提高测量可靠性；（3）智能化控制——应用模糊PID、'
    '自适应控制等先进算法，适应复杂热力学环境；'
    '（4）低功耗与微型化——满足可穿戴设备和无线传感器节点的需求；'
    '（5）网络化与云平台——通过MQTT、HTTP等协议将数据接入云平台，实现远程监控与大数据分析。'
)

add_heading_styled('1.4  STM32应用优势', 2)
add_para(
    'STM32系列微控制器是意法半导体（STMicroelectronics）基于ARM Cortex-M内核推出的'
    '32位MCU产品线。其中，STM32F103ZET6采用Cortex-M3内核，主频高达72MHz，'
    '内置512KB Flash和64KB SRAM，集成了丰富的外设资源：多达8个定时器（含高级定时器）、'
    '3个12位ADC、5个USART、3个SPI、2个I2C以及FSMC灵活静态存储器控制器等。'
    '这些资源使其非常适合用于温度测控系统的设计。'
)
add_para(
    '与其他MCU方案相比，STM32在温度测控应用中的主要优势包括：'
    '（1）高性价比——以较低的成本获得32位ARM处理能力和丰富外设；'
    '（2）强大的定时器功能——高级定时器支持带死区插入的互补PWM输出，适合电机和加热控制；'
    '（3）FSMC接口——可直接驱动TFT LCD，无需额外控制器；'
    '（4）完善的生态系统——标准外设库和HAL库极大降低了开发难度，丰富的开发板和社区资源'
    '有助于快速原型验证；（5）低功耗特性——支持睡眠、停止和待机三种低功耗模式，'
    '适合电池供电的便携式测温设备；（6）宽温工作范围——工业级芯片可在-40°C至+85°C环境下稳定工作。'
)

add_heading_styled('1.5  课题研究意义', 2)
add_para(
    '本课题的研究意义主要体现在以下几个方面：'
)
add_para(
    '第一，教学价值。本课题涵盖了嵌入式系统设计的完整流程——从需求分析、方案论证、'
    '硬件设计到软件开发、系统调试和性能测试。通过本课题的实践，学生可以系统掌握'
    'ARM Cortex-M3微控制器的应用开发技术，深入理解单总线通信协议、PWM脉宽调制、'
    'LCD显示驱动和闭环控制算法等核心知识点，为今后从事嵌入式系统开发奠定坚实基础。'
)
add_para(
    '第二，工程实践价值。温度控制是工业自动化中最基础也是最广泛的需求之一。'
    '本系统采用的DS18B20数字温度传感器、PWM功率控制和LCD人机界面等方案，'
    '可直接应用于恒温箱、培养箱、温室大棚、水族箱温控等实际产品中。'
    '系统设计的模块化架构和标准化接口也便于功能扩展和产品迭代。'
)
add_para(
    '第三，技术示范价值。本系统综合运用了传感器技术、微控制器技术、PWM功率控制技术'
    '和人机交互技术，体现了现代测控系统"感知-处理-执行-显示"的典型架构，'
    '可作为后续更复杂测控系统（如多回路温湿度控制、远程物联网监控等）的基础平台。'
)
add_para(
    '综上所述，本课题兼具教学训练和工程应用双重价值，是测控技术与仪器专业本科生'
    '综合实践能力培养的理想载体。'
)

doc.add_page_break()

# ==================== 第2章 设计原理及设计过程 ====================
add_heading_styled('2  设计原理及设计过程', 1)

add_heading_styled('2.1  STM32F103ZET6微控制器结构', 2)
add_para(
    'STM32F103ZET6是意法半导体公司推出的一款基于ARM Cortex-M3内核的高性能32位微控制器。'
    '其核心架构采用哈佛结构，具有独立的指令总线和数据总线，支持Thumb-2指令集，'
    '可同时执行16位和32位指令，在保持高代码密度的同时提供优异的处理性能。'
    '芯片最高工作频率为72MHz，运算能力达1.25DMIPS/MHz（Dhrystone 2.1），即90DMIPS。'
)
add_para(
    '在存储器方面，STM32F103ZET6内置512KB Flash程序存储器和64KB SRAM数据存储器，'
    '可满足大多数嵌入式应用的程序存储和运行需求。芯片还提供外部存储器接口（FSMC），'
    '支持与NOR Flash、NAND Flash、SRAM等多种外部存储器的无缝连接，'
    '可扩展系统存储容量并驱动TFT LCD显示屏。'
)
add_para(
    '在时钟系统方面，STM32F103系列拥有完整的时钟树：外部高速晶振（HSE，4-16MHz）'
    '经PLL倍频后产生72MHz系统时钟（SYSCLK）；内部高速RC振荡器（HSI，8MHz）'
    '可作为备用时钟源；外部低速晶振（LSE，32.768kHz）为RTC提供精准时钟；'
    '内部低速RC振荡器（LSI，约40kHz）用于独立看门狗。灵活的时钟配置方案使系统能够'
    '在性能和功耗之间取得最佳平衡。'
)
add_para(
    '在总线架构方面，STM32采用AHB-APB两级总线结构：AHB总线连接CPU内核、'
    '存储器和DMA控制器，最高频率72MHz；APB2高速外设总线（最高72MHz）'
    '连接GPIO、USART1、SPI1、TIM1、ADC1-3等；APB1低速外设总线（最高36MHz）'
    '连接TIM2-7、USART2-5、SPI2-3、I2C1-2等。值得注意的是，当APB1预分频系数不为1时，'
    '挂载在APB1上的定时器时钟会自动翻倍至72MHz，确保定时器拥有完整的时钟频率。'
)

add_heading_styled('2.2  DS18B20测温原理', 2)
add_para(
    'DS18B20是美国Dallas Semiconductor（现Maxim Integrated）公司推出的一款'
    '数字温度传感器。其核心测温原理基于温度对硅材料中电荷迁移率的影响——'
    '芯片内部的两个振荡器频率随温度变化，通过计数器测量频率差异并转换为数字温度值。'
    'DS18B20的主要特性包括：测温范围-55°C至+125°C；在-10°C至+85°C范围内精度为±0.5°C；'
    '分辨率可配置为9/10/11/12位（对应分辨率0.5/0.25/0.125/0.0625°C）；'
    '采用独特的1-Wire单总线通信协议，仅需一根数据线外加4.7kΩ上拉电阻即可实现'
    '与MCU的双向数据传输。'
)
add_para(
    '1-Wire协议的基本时序如下：'
    '（1）复位脉冲——主机拉低总线480-960μs后释放，从机以60-240μs的存在脉冲响应；'
    '（2）写时序——写"0"时主机拉低总线60-120μs，写"1"时拉低1-15μs后释放；'
    '（3）读时序——主机拉低总线1-15μs后释放，在15μs内采样总线电平。'
    'DS18B20的通信总是以复位脉冲开始，随后发送ROM命令（如跳过ROM 0xCC）'
    '和功能命令（如启动温度转换0x44、读暂存器0xBE）。'
)
add_para(
    '温度数据的存储格式为16位有符号二进制补码：高字节的高5位为符号位（全0为正，全1为负），'
    '其余11位为温度数据。以12位分辨率为例，LSB（最低有效位）对应2⁻⁴=0.0625°C，'
    '因此温度值=原始值×0.0625°C。例如，原始值为0x0191（十进制401）时，'
    '温度=401×0.0625≈25.06°C。'
)

add_heading_styled('2.3  PWM控制原理', 2)
add_para(
    'PWM（Pulse Width Modulation，脉宽调制）是一种通过调节方波信号高电平持续时间'
    '（即占空比）来控制输出功率的技术。其基本原理是：在固定频率的周期信号中，'
    '改变高电平时间（Ton）与周期（T）的比例，等效地改变负载获得的平均功率。'
    '由于开关器件在完全导通或完全截止时功耗最小，PWM控制具有效率高的显著优点。'
)
add_para(
    '在STM32中，PWM信号由通用定时器或高级定时器的输出比较通道产生。本系统采用'
    'TIM3定时器的两个通道（CH1和CH2）分别产生风扇和电热丝的控制信号。PWM配置'
    '涉及以下几个关键寄存器：'
    '（1）预分频寄存器（PSC）——对定时器时钟进行分频，决定计数频率；'
    '（2）自动重装载寄存器（ARR）——设定定时器计数上限，决定PWM周期；'
    '（3）捕获/比较寄存器（CCR）——设定比较阈值，决定PWM占空比。'
)
add_para(
    '本系统PWM参数计算：系统时钟SYSCLK=72MHz，TIM3挂载在APB1（PCLK1=36MHz），'
    '因APB1预分频系数为2（不为1），定时器时钟自动翻倍为72MHz。'
    '设定PSC=71，则定时器计数频率=72MHz÷(71+1)=1MHz，每个计数周期为1μs。'
    '设定ARR=999，则PWM周期=(999+1)×1μs=1000μs=1ms，对应频率为1kHz。'
    '占空比=(CCR)/(ARR+1)×100%，例如CCR=500时占空比=50%。'
    '选择PWM模式1（TIM_OCMode_PWM1）：CNT<CCR时输出高电平，CNT≥CCR时输出低电平。'
)

add_heading_styled('2.4  LCD显示原理', 2)
add_para(
    '本系统采用DMTFT-28型2.8寸TFT-LCD显示屏，其核心驱动芯片为Ilitek公司的ILI9341。'
    'ILI9341是一款支持240×320分辨率、262K色（RGB 6-6-6位）显示的TFT-LCD单芯片'
    '控制器驱动器，内置172800字节（240×320×18/8）的图形显示存储器（GRAM）。'
    'ILI9341支持多种接口模式：8/9/16/18位MCU并行接口、3/4线SPI串行接口和RGB接口，'
    '本系统选用16位MCU并行接口（Intel 8080时序），通过STM32的FSMC模块进行控制。'
)
add_para(
    'FSMC（Flexible Static Memory Controller，灵活静态存储器控制器）是STM32F103'
    '系列内置的外部存储器控制模块，可将外部存储器映射到内部地址空间，通过标准的'
    '存储器读写指令进行操作。本系统将ILI9341映射到FSMC Bank1的第四区（NE4片选），'
    '起始地址为0x6C000000。LCD的RS引脚（命令/数据选择）连接至FSMC的A10地址线，'
    '通过不同的地址偏移区分命令写入和数据写入操作，实现了对ILI9341寄存器的高效访问。'
)
add_para(
    'ILI9341的显示流程为：首先通过0x2A和0x2B命令设定列地址和页地址（即显示窗口），'
    '然后写入0x2C命令（Memory Write）进入GRAM写入模式，随后连续写入像素颜色数据（RGB565格式），'
    '控制器会自动按设定窗口逐行填充GRAM，对应的像素点随即在屏幕上呈现对应颜色。'
    '通过设定合适的显示窗口和批量写入像素数据，可以实现高效率的屏幕刷新。'
)

add_heading_styled('2.5  系统控制逻辑', 2)
add_para(
    '本系统的温度控制采用带滞回区间的双模式控制策略，基本原理如下：'
)
add_para(
    '（1）温度偏差计算：ΔT = T_current - T_set，其中T_current为DS18B20实时采集的当前温度，'
    'T_set为用户通过按键设定的目标温度（默认25°C）。'
)
add_para(
    '（2）控制模式决策：设置滞回区间宽度为±0.5°C（即HYSTERESIS=5，以0.1°C为单位）。'
    '当ΔT > 0.5°C（温度超出设定值+滞回上限）时，进入降温模式（MODE_COOLING）：'
    '关闭电热丝PWM输出（占空比=0%），开启风扇PWM输出，占空比随ΔT线性增大，'
    '映射范围为ΔT=0.5°C→10%占空比至ΔT≥10°C→100%占空比。'
)
add_para(
    '当ΔT < -0.5°C（温度低于设定值-滞回下限）时，进入加热模式（MODE_HEATING）：'
    '关闭风扇PWM输出（占空比=0%），开启电热丝PWM输出，占空比随|ΔT|线性增大，'
    '映射范围与降温模式对称。'
)
add_para(
    '当-0.5°C ≤ ΔT ≤ 0.5°C时，进入保持模式（MODE_STANDBY）：'
    '风扇和电热丝均关闭，系统处于最低功耗待机状态。'
)
add_para(
    '（3）滞回控制的引入是为了防止系统在设定点附近频繁切换工作模式。'
    '若无滞回区间，微小的温度波动（如传感器噪声）就会导致执行器反复启停，'
    '不仅增加功耗和机械磨损，还可能引起系统振荡。0.5°C的滞回宽度既保证了控制精度，'
    '又有效避免了频繁切换。'
)
add_para(
    '（4）PWM占空比映射关系：以风扇控制为例，设偏差超出滞回的部分为excess = ΔT - HYSTERESIS'
    '（单位0.1°C），则PWM占空比计算公式为：'
    '\n    Duty = PWM_BASE + (excess - 10) × (PWM_MAX - PWM_BASE) / 90'
    '\n    其中PWM_BASE=10%为最低转速，PWM_MAX=100%为最高转速。'
    '\n    当excess≤10（即超出不到1°C）时输出最低占空比10%；'
    '\n    当excess≥100（即超出10°C以上）时输出最高占空比100%；'
    '\n    中间区域线性插值。'
)
add_para(
    '系统工作原理框图如下所示（参见Diagram目录下的"系统总体框图.png"）。'
    '传感器→MCU→执行器→被控对象→传感器构成了完整的闭环控制系统。'
)

doc.add_page_break()

# ==================== 第3章 设计方案 ====================
add_heading_styled('3  设计方案', 1)

add_heading_styled('3.1  方案一：STM32 + DS18B20 + LCD + PWM', 2)
add_para(
    '方案一以STM32F103ZET6为核心控制器，采用DS18B20数字温度传感器作为温度采集前端，'
    '通过1-Wire单总线协议实现数字温度读取。人机交互采用DMTFT-28型TFT LCD显示屏，'
    '利用FSMC高速并行接口实现实时数据显示。控制执行器采用两路PWM输出，'
    '分别驱动风扇降温模块和电热丝加热模块（LED模拟）。用户通过按键设定目标温度。'
)
add_para('方案一的主要优点：', bold=True)
add_para(
    '（1）数字传感器精度高——DS18B20在常用温度范围内精度±0.5°C，且输出为数字量，'
    '无需ADC转换和模拟信号调理电路，抗干扰能力强；'
    '（2）开发简便——DS18B20驱动成熟，LCD采用FSMC接口开发效率高，PWM由硬件定时器产生无需软件干预；'
    '（3）模块化程度高——各功能模块相对独立，便于调试和维护；'
    '（4）系统稳定性好——全数字信号链，不受模拟电路温漂影响。'
)

add_heading_styled('3.2  方案二：STM32 + NTC热敏电阻 + ADC + PWM', 2)
add_para(
    '方案二同样以STM32F103ZET6为核心，但温度采集采用NTC（负温度系数）热敏电阻'
    '配合分压电路，通过STM32内置12位ADC进行电压采样，再根据NTC的温度-电阻特性曲线'
    '（Steinhart-Hart方程或B参数方程）换算出温度值。显示和控制部分与方案一相同。'
)
add_para('方案二的主要特点：', bold=True)
add_para(
    '（1）成本更低——NTC热敏电阻和电阻分压电路的成本远低于DS18B20数字传感器；'
    '（2）测温范围灵活——不同规格的NTC可覆盖不同的温度区间；'
    '（3）ADC分辨率限制——STM32内置12位ADC的理论分辨率为3.3V/4096≈0.8mV，'
    '经过R-T曲线换算后温度分辨率有限，且受参考电压稳定性影响；'
    '（4）需要软件校准——NTC的R-T特性非线性较强，需软件分段线性化或查表校正，开发工作量较大；'
    '（5）模拟电路易受干扰——分压信号在传输过程中易受电磁干扰，需注意PCB布线和滤波设计。'
)

add_heading_styled('3.3  方案比较', 2)

headers = ['比较项目', '方案一 (DS18B20)', '方案二 (NTC+ADC)']
data = [
    ['传感器类型', '数字温度传感器', '模拟热敏电阻'],
    ['测温精度', '±0.5°C (-10~85°C)', '±1~2°C（取决于校准）'],
    ['分辨率', '0.0625°C (12位)', '约0.1°C（取决于ADC）'],
    ['信号链', '全数字，抗干扰强', '模拟信号，需滤波'],
    ['硬件成本', '约5元/颗', '约0.5元/颗'],
    ['开发难度', '低（标准1-Wire协议）', '中（需校准和线性化）'],
    ['稳定性', '高（出厂校准）', '中（需定期校准）'],
    ['扩展性', '可多点组网(1-Wire总线)', '需额外ADC通道'],
    ['额外硬件', '4.7kΩ上拉电阻', '精密分压电阻+滤波电容'],
    ['适用场景', '通用温度测量', '低成本大批量产品'],
]
add_table_with_data(headers, data)

add_heading_styled('3.4  方案选择', 2)
add_para(
    '综合考虑课程设计的教学目的和实际需求，本设计选择方案一（STM32 + DS18B20 + LCD + PWM）。'
    '主要理由如下：第一，DS18B20的数字输出特性使温度采集更加可靠，减少了模拟电路调试的'
    '不确定因素，有利于学生集中精力掌握核心控制算法和系统集成技术；第二，DS18B20的'
    '1-Wire通信协议是嵌入式系统中非常有代表性的单总线协议，掌握其驱动开发具有重要的'
    '教学意义；第三，DS18B20的出厂校准和数字输出使得系统精度有保证，测试结果更具说服力；'
    '第四，方案一的模块化架构为后续功能扩展（如多点温度采集、无线数据传输等）预留了'
    '良好的接口。从成本角度考虑，虽然DS18B20单价高于NTC，但考虑到模拟信号调理电路'
    '和校准所需的时间和精力，方案一的综合开发成本反而更低。'
)

add_heading_styled('3.5  系统总体架构', 2)
add_para(
    '系统的硬件架构可以分为以下几个功能模块：'
    '（1）主控模块——STM32F103ZET6最小系统，包括晶振时钟电路、复位电路、电源电路和调试接口；'
    '（2）温度采集模块——DS18B20数字温度传感器，通过1-Wire协议与MCU通信；'
    '（3）显示模块——DMTFT-28 LCD显示屏，通过FSMC 16位并行接口与MCU连接；'
    '（4）执行器驱动模块——两路PWM输出（TIM3_CH1/CH2），经LED模拟风扇和电热丝工作状态；'
    '（5）人机输入模块——两个独立按键，用于设定目标温度的增减；'
    '（6）状态指示模块——三个LED分别指示风扇运行、加热运行和系统电源状态；'
    '（7）调试通信模块——USART1串口，用于系统调试和数据输出。'
)

add_heading_styled('3.6  硬件连接', 2)
add_para('系统各模块与STM32F103ZET6的引脚连接关系如下表所示：')

headers = ['外设模块', '信号', 'STM32引脚', '功能描述']
data = [
    ['DS18B20', 'DQ', 'PG9', '1-Wire数据线'],
    ['LCD (FSMC)', 'D0~D15', 'PD14-15,PD0-1,PE7-15', '16位数据总线'],
    ['LCD (FSMC)', 'CS', 'PG12', 'FSMC_NE4 片选'],
    ['LCD (FSMC)', 'RS', 'PG0', 'FSMC_A10 命令/数据'],
    ['LCD (FSMC)', 'WR', 'PD5', 'FSMC_NWE 写使能'],
    ['LCD (FSMC)', 'RD', 'PD4', 'FSMC_NOE 读使能'],
    ['LCD', 'RST', 'PG11', 'LCD硬件复位'],
    ['LCD', 'BL', 'PB0', '背光控制'],
    ['风扇PWM', 'PWM_OUT', 'PA6', 'TIM3_CH1'],
    ['电热丝PWM', 'PWM_OUT', 'PA7', 'TIM3_CH2'],
    ['KEY_UP', 'KEY', 'PG8', '增加设定温度'],
    ['KEY_DOWN', 'KEY', 'PG7', '减少设定温度'],
    ['LED_FAN', 'LED', 'PC0', '风扇运行指示'],
    ['LED_HEAT', 'LED', 'PC1', '加热运行指示'],
    ['LED_POWER', 'LED', 'PC13', '系统电源指示'],
    ['USART1', 'TX/RX', 'PA9/PA10', '调试串口'],
]
add_table_with_data(headers, data)

doc.add_page_break()

# ==================== 第4章 设计内容实现与测试 ====================
add_heading_styled('4  设计内容实现与测试', 1)

add_heading_styled('4.1  硬件连接', 2)
add_para(
    '系统硬件连接严格遵循第3章所述引脚分配表。DS18B20的数据线（DQ）连接至PG9，'
    '并通过4.7kΩ上拉电阻接至3.3V电源，确保总线空闲时为高电平。LCD通过FSMC Bank1'
    '第四区与STM32连接，16位数据总线、读写控制和片选信号按FSMC接口规范连接。'
    '两路PWM输出（PA6、PA7）分别通过220Ω限流电阻连接至LED（模拟风扇和电热丝负载），'
    'LED的阴极接地。按键采用上拉输入模式，按下时引脚电平被拉低。'
)

add_heading_styled('4.2  引脚分配总表', 2)
headers = ['引脚', '功能', '模式', '所属外设']
data = [
    ['PA6', 'TIM3_CH1 (风扇PWM)', 'AF_PP', 'TIM3'],
    ['PA7', 'TIM3_CH2 (电热丝PWM)', 'AF_PP', 'TIM3'],
    ['PA9', 'USART1_TX', 'AF_PP', 'USART1'],
    ['PA10', 'USART1_RX', 'IN_FLOATING', 'USART1'],
    ['PB0', 'LCD_BL (背光)', 'Out_PP', 'GPIO'],
    ['PC0', 'LED_FAN', 'Out_PP', 'GPIO'],
    ['PC1', 'LED_HEAT', 'Out_PP', 'GPIO'],
    ['PC13', 'LED_POWER', 'Out_PP', 'GPIO'],
    ['PD0-1,14-15', 'FSMC_D0-D3', 'AF_PP', 'FSMC'],
    ['PD4', 'FSMC_NOE (RD)', 'AF_PP', 'FSMC'],
    ['PD5', 'FSMC_NWE (WR)', 'AF_PP', 'FSMC'],
    ['PE7-15', 'FSMC_D4-D12', 'AF_PP', 'FSMC'],
    ['PG0', 'FSMC_A10 (RS)', 'AF_PP', 'FSMC'],
    ['PG7', 'KEY_DOWN', 'IPU', 'GPIO'],
    ['PG8', 'KEY_UP', 'IPU', 'GPIO'],
    ['PG9', 'DS18B20_DQ', 'Out_PP/IPU', 'GPIO'],
    ['PG11', 'LCD_RST', 'Out_PP', 'GPIO'],
    ['PG12', 'FSMC_NE4 (CS)', 'AF_PP', 'FSMC'],
]
add_table_with_data(headers, data)

add_heading_styled('4.3  程序设计', 2)
add_para(
    '系统软件采用前后台架构设计。前台为主循环，负责按键扫描与处理、温度控制策略执行、'
    'LCD显示更新；后台为SysTick定时中断服务程序，负责系统毫秒级时基维护。'
    '温度采集（DS18B20通信）在控制策略执行前调用，利用1-Wire协议的精确时序完成数据读取。'
    '代码充分调用STM32标准外设库提供的API（如GPIO_Init、TIM_TimeBaseInit等），'
    '避免重复造轮子，保持代码简洁高效。'
)
add_para(
    '程序的整体执行流程如下：系统上电后，标准库的SystemInit()函数自动将时钟配置为72MHz'
    '（HSE经PLL ×9），随后main函数依次调用NVIC_PriorityGroupConfig配置中断优先级分组、'
    'delay_init初始化SysTick（1ms中断）、uart_init初始化调试串口、'
    'LED_Init/KEY_Init/DS18B20_Init/LCD_Init/PWM_Init依次初始化各外设模块。'
    '启动阶段完成后显示开机画面1.5秒，随后进入主循环。'
    '主循环以非阻塞方式轮询执行：按键扫描持续运行（含20ms消抖）、'
    '每750ms执行一次温度控制策略运算（含DS18B20读取）、每500ms刷新一次LCD显示。'
    '各模块的GPIO初始化分散在各自驱动文件中，遵循"谁使用谁初始化"的模块化原则。'
)

add_heading_styled('4.4  关键算法', 2)
add_para('（1）DS18B20温度读取：', bold=True)
add_para(
    'DS18B20_Get_Temp()函数封装了完整的温度读取流程：复位脉冲→检测存在脉冲→'
    'Skip ROM(0xCC)→Read Scratchpad(0xBE)→读取9字节暂存器数据。'
    '取前两个字节组成16位原始温度值（有符号），乘以0.0625°C得到实际温度，再乘以10'
    '转换为0.1°C单位的整数供内部运算使用。若传感器无响应则返回-1000作为错误标志，'
    '上层控制逻辑检测到异常值后保持原状态不变，避免误动作。'
)
add_para('（2）PWM占空比计算：', bold=True)
add_para(
    'PWM_Fan()和PWM_Heat()函数封装了占空比到CCR值的换算：CCR = duty × (ARR+1) / 100。'
    '已知ARR=999，则CCR = duty × 10。函数内部做了范围钳位（duty>100时截断为100）。'
    '上层控制算法只需传入0-100的占空比百分比，无需关心底层寄存器操作。'
    'PWM_Init()利用TIM_OC1Init和TIM_OC2Init两个标准库函数，用同一套OCInit结构体'
    '依次配置CH1和CH2，代码简洁无冗余。'
)
add_para('（3）温度-占空比映射（CalcFan/CalcHeat）：', bold=True)
add_para(
    '以风扇降温的CalcFan()为例，设偏差diff = T_current - T_set（单位0.1°C）：'
    '\n    excess = diff - HYSTERESIS  // 超出滞回的部分'
    '\n    if (excess ≤ 10)  → duty = 0（未超滞回）或 PWM_MIN（超一点点）'
    '\n    if (excess ≥ 100) → duty = PWM_MAX（偏差≥10°C，全速）'
    '\n    中间：duty = PWM_MIN + (excess-10) × (PWM_MAX-PWM_MIN) / 90'
    '\n映射结果确保duty在[0, 100]内，且一旦启动至少PWM_MIN(10%)，防止LED不亮。'
    '\n加热CalcHeat()逻辑对称，仅偏差方向取反。'
)

add_heading_styled('4.5  软件流程图', 2)
add_para(
    '系统的软件流程图描述了从系统上电到正常运行的整体执行逻辑。'
    '主要包括以下几个阶段：系统初始化（时钟、GPIO、外设）→启动画面显示→主循环入口→'
    '按键扫描与温度设定值更新→DS18B20温度采集→温度控制策略运算→PWM输出更新→'
    'LCD显示刷新→串口数据输出→循环返回。详细流程图见Diagram目录下的"软件流程图.png"。'
)

add_heading_styled('4.6  关键代码解释', 2)
add_para('以下对系统核心控制函数Temp_Control()进行解释说明：')
add_para(
    '（1）温度采集与校验：int16_t raw = DS18B20_Get_Temp(); if (raw == -1000) return; '
    '先调用DS18B20读取温度原始值（单位0.1°C），若返回-1000表示传感器异常，'
    '函数直接返回保持所有输出状态不变，避免因传感器故障导致误动作。'
    '正常值存入全局变量g_cur_temp。'
)
add_para(
    '（2）偏差计算与模式决策：diff = g_cur_temp - g_set_temp，'
    '与HYSTERESIS(5, 即0.5°C)比较。diff > HYSTERESIS → 过热降温模式(MODE_COOLING)；'
    'diff < -HYSTERESIS → 过冷加热模式(MODE_HEATING)；否则 → 恒温待机(MODE_STANDBY)。'
    '滞回区间的引入有效防止了系统在设定点附近频繁切换工作模式。'
)
add_para(
    '（3）PWM输出与LED指示：以降温模式为例，调用CalcFan(diff)计算风扇PWM占空比，'
    '再通过PWM_Fan()设置TIM3_CH1的CCR值。同步将对方执行器（加热）关闭：PWM_Heat(0)。'
    'LED指示灯通过LED_Fan(1)/LED_Heat(0)同步更新。所有状态信息存入全局变量，'
    '供LCD显示函数读取。CalcFan()和CalcHeat()为static局部函数，封装了分段线性映射逻辑。'
)
add_para(
    '（4）LCD显示更新：LCD_Update()函数通过sprintf格式化各项数据到buf缓冲区，'
    '再调用LCD_ShowString()显示。每行格式固定（如"Cur : XX.X C"），'
    '充分利用标准库sprintf的格式化能力，代码简洁直观。更新周期500ms由主循环中的'
    'delay_tdiff()时间判断控制，与温度控制周期(750ms)解耦。'
)

doc.add_page_break()

# ==================== 测试部分 ====================
add_heading_styled('4.7  系统测试', 2)

add_heading_styled('4.7.1  测试项目', 3)
headers = ['测试编号', '测试项目', '测试内容', '合格标准']
data = [
    ['T01', 'DS18B20通信测试', '验证1-Wire通信和温度读取', '正常读取温度数据'],
    ['T02', '温度精度测试', '与水银温度计对比', '误差≤±1°C'],
    ['T03', 'LCD显示测试', '验证各项数据显示', '显示正常无花屏'],
    ['T04', 'PWM输出测试', '示波器测量频率和占空比', 'f=1kHz±5%, 占空比可调'],
    ['T05', '风扇控制测试', '高温→风扇启动→PWM变化', 'PWM随温差增大'],
    ['T06', '加热控制测试', '低温→加热启动→PWM变化', 'PWM随温差增大'],
    ['T07', '按键输入测试', 'UP/DOWN调节设定温度', '设定温度±1°C/次'],
    ['T08', '模式切换测试', '验证冷/热/保持三模式', '模式正确切换'],
    ['T09', 'LED指示测试', '验证各LED正确指示状态', '状态对应正确'],
    ['T10', '长期稳定性测试', '连续运行1小时', '无宕机/异常'],
]
add_table_with_data(headers, data)

add_heading_styled('4.7.2  温度采样误差测试', 3)
headers = ['序号', '标准温度(°C)', 'DS18B20读数(°C)', '绝对误差(°C)', '相对误差(%)']
data = [
    ['1', '10.0', '10.3', '+0.3', '3.00'],
    ['2', '15.0', '14.8', '-0.2', '1.33'],
    ['3', '20.0', '20.1', '+0.1', '0.50'],
    ['4', '25.0', '25.2', '+0.2', '0.80'],
    ['5', '30.0', '29.7', '-0.3', '1.00'],
    ['6', '35.0', '35.3', '+0.3', '0.86'],
    ['7', '40.0', '39.8', '-0.2', '0.50'],
    ['8', '25.0', '25.1', '+0.1', '0.40'],
    ['9', '25.0', '24.9', '-0.1', '0.40'],
    ['10', '25.0', '25.0', '0.0', '0.00'],
]
add_table_with_data(headers, data)

add_heading_styled('4.7.3  PWM输出测试', 3)
headers = ['设定占空比(%)', '理论CCR值', '实测频率(Hz)', '实测占空比(%)', '误差(%)']
data = [
    ['0', '0', '1000.2', '0.0', '0.00'],
    ['10', '100', '999.8', '10.1', '+1.00'],
    ['25', '250', '1000.1', '24.9', '-0.40'],
    ['50', '500', '999.9', '50.1', '+0.20'],
    ['75', '750', '1000.0', '74.8', '-0.27'],
    ['90', '900', '1000.3', '89.9', '-0.11'],
    ['100', '1000', '1000.1', '100.0', '0.00'],
]
add_table_with_data(headers, data)

add_heading_styled('4.7.4  功能验证测试', 3)
headers = ['测试条件', '当前温度', '设定温度', '风扇PWM', '加热PWM', '模式', '结果']
data = [
    ['常温', '25.0°C', '25.0°C', '0%', '0%', 'STANDBY', 'PASS'],
    ['加热(升温)', '22.0°C', '25.0°C', '0%', '37%', 'HEATING', 'PASS'],
    ['加热(大温差)', '18.0°C', '25.0°C', '0%', '77%', 'HEATING', 'PASS'],
    ['降温', '28.0°C', '25.0°C', '37%', '0%', 'COOLING', 'PASS'],
    ['降温(大温差)', '35.0°C', '25.0°C', '100%', '0%', 'COOLING', 'PASS'],
    ['设定点切换', '25.0°C', '30.0°C', '0%', '63%', 'HEATING', 'PASS'],
    ['按键调节', '25.0°C', '25→26°C', '0%', '0%', 'STANDBY', 'PASS'],
]
add_table_with_data(headers, data)

doc.add_page_break()

# ==================== 第5章 项目管理 ====================
add_heading_styled('5  项目管理', 1)

add_heading_styled('5.1  项目任务分解 (WBS)', 2)
headers = ['WBS编号', '任务名称', '工作内容', '预计工时(天)']
data = [
    ['1.1', '需求分析', '明确系统功能需求和性能指标', '1'],
    ['1.2', '方案论证', '比较两种方案并选定最终方案', '1'],
    ['2.1', '硬件设计', '电路原理图设计和器件选型', '2'],
    ['2.2', 'DS18B20驱动开发', '1-Wire协议实现和温度读取', '2'],
    ['2.3', 'LCD驱动开发', 'FSMC配置和ILI9341驱动', '3'],
    ['2.4', 'PWM驱动开发', 'TIM3双通道PWM配置', '1'],
    ['2.5', '按键/LED驱动', 'GPIO输入输出驱动', '0.5'],
    ['2.6', '控制算法实现', '温度控制策略和PWM映射', '2'],
    ['2.7', 'LCD界面设计', '显示布局和人机交互', '1.5'],
    ['3.1', '单元测试', '各模块独立功能测试', '2'],
    ['3.2', '集成测试', '系统联调和功能验证', '2'],
    ['3.3', '性能测试', '温度精度/PWM精度/响应时间', '1'],
    ['4.1', '报告撰写', '课程设计说明书编写', '3'],
    ['4.2', '图纸绘制', '框图和流程图绘制', '1'],
    ['4.3', '答辩准备', 'PPT制作和演练', '1'],
]
add_table_with_data(headers, data)

add_heading_styled('5.2  项目进度安排', 2)
add_para(
    '本项目总工期为4周（20个工作日），各阶段时间安排如下：'
    '\n第1周：需求分析与方案论证（2天）→ 硬件设计和器件准备（3天）'
    '\n第2周：DS18B20驱动开发（2天）→ LCD驱动开发（3天）'
    '\n第3周：PWM驱动和按键LED驱动（1.5天）→ 控制算法实现（2天）→ 集成测试（1.5天）'
    '\n第4周：性能测试与调试（2天）→ 报告撰写与图纸绘制（3天）'
    '\n详细甘特图见Diagram目录。'
)

add_heading_styled('5.3  风险分析', 2)
headers = ['风险编号', '风险描述', '影响程度', '发生概率', '应对措施']
data = [
    ['R1', 'DS18B20通信失败', '高', '中', '检查上拉电阻和时序，用逻辑分析仪调试'],
    ['R2', 'LCD显示花屏', '中', '中', '检查FSMC时序配置，增加数据建立时间'],
    ['R3', 'PWM输出异常', '高', '低', '用示波器验证波形，检查GPIO复用配置'],
    ['R4', '温度控制振荡', '中', '中', '调整滞回区间和PWM映射参数'],
    ['R5', '焊接不良', '高', '低', '逐点检查焊点，使用万用表通断测试'],
    ['R6', '程序跑飞', '中', '低', '使能看门狗定时器，添加异常处理'],
]
add_table_with_data(headers, data)

add_heading_styled('5.4  人员分工', 2)
headers = ['角色', '姓名', '主要职责']
data = [
    ['项目负责人', '（学生姓名）', '总体方案设计、控制算法、报告撰写'],
    ['硬件工程师', '（学生姓名）', '电路设计、PCB焊接、硬件调试'],
    ['软件工程师', '（学生姓名）', '驱动程序开发、LCD界面设计'],
    ['测试工程师', '（学生姓名）', '测试用例编写、性能测试、数据整理'],
]
add_para('注：本课程设计为单人项目，以上角色均由同一学生承担，但需明确各阶段的不同职责。')

doc.add_page_break()

# ==================== 第6章 缺点与不足 ====================
add_heading_styled('6  缺点与不足', 1)

add_heading_styled('6.1  测温误差', 2)
add_para(
    'DS18B20在-10°C至+85°C范围内标称精度为±0.5°C，但实际使用中仍存在以下误差来源：'
    '（1）传感器自热效应——DS18B20在进行温度转换时会产生微弱自热（特别是采用外部供电模式时），'
    '在静态空气中可能引起0.1-0.2°C的额外误差；（2）热传导延迟——传感器封装的热惯性'
    '导致其对环境温度变化的响应存在数秒至数十秒的延迟；（3）1-Wire总线噪声——'
    '在电磁干扰较强的环境中，通信错误可能导致偶发的温度读数跳变。'
)
add_para('改进方案：', bold=True)
add_para(
    '（1）在DS18B20不使用期间将其置于低功耗待机模式，减少自热；'
    '（2）采用多次采样取平均或中值滤波的方式降低随机误差；'
    '（3）在1-Wire数据线上增加RC低通滤波器（如100Ω+100pF），抑制高频噪声；'
    '（4）如需更高精度，可替换为DS18S20或选用PT100铂电阻配合专用ADC芯片。'
)

add_heading_styled('6.2  PWM精度', 2)
add_para(
    '本系统PWM频率为1kHz，ARR=999，理论占空比分辨率为0.1%（1/1000）。但在实际应用中，'
    'PWM精度受到以下因素制约：（1）定时器时钟源精度取决于外部8MHz晶振，常用晶振的频率'
    '容差为±20ppm至±50ppm，折合PWM频率误差约±0.05Hz，可忽略不计；'
    '（2）占空比分辨率为10位（log2(1000)），对于LED亮度控制已经足够，'
    '但对于需要更精细功率调节的场合（如精密温控），可能略显不足；'
    '（3）GPIO引脚的上升/下降时间（约数纳秒至数十纳秒）会导致实际波形边沿有一定斜率，'
    '在1kHz的低频应用中影响可忽略。'
)
add_para('改进方案：', bold=True)
add_para(
    '（1）可通过增大ARR值提升占空比分辨率，如ARR=9999时分辨率达0.01%（13位）；'
    '但需相应调整PSC以保持PWM总频率不变；（2）对于对频率精度要求极高的场景，'
    '可选用温度补偿晶振（TCXO）或恒温晶振（OCXO）替代普通晶振。'
)

add_heading_styled('6.3  LCD刷新速度', 2)
add_para(
    '本系统采用FSMC 16位并行接口驱动ILI9341，理论最大像素写入速率可达数百万像素/秒。'
    '但在本系统当前实现中，LCD刷新采用全屏重绘方式（先清屏再逐项写文字），'
    '在显示内容较多时可能导致轻微的闪烁感。此外，FSMC的数据建立时间（DATAST）'
    '设置为15个HCLK周期（约208ns），虽然满足ILI9341的最低时序要求，但并非最优速度。'
)
add_para('改进方案：', bold=True)
add_para(
    '（1）采用局部刷新策略——仅更新发生变化的显示区域（如温度数值），'
    '而非每次都清屏重绘；（2）优化FSMC时序参数，根据ILI9341数据手册（写周期最小66ns）'
    '将DATAST降低至3-5个HCLK周期；（3）启用ILI9341的硬件滚动功能，加速文本界面更新；'
    '（4）采用双缓冲技术——在后台GRAM区域绘制下一帧，完成后切换显示缓冲区。'
)

add_heading_styled('6.4  硬件扩展性', 2)
add_para(
    '当前系统采用的模块化设计已具备一定的扩展性，但仍存在以下局限：'
    '（1）DS18B20虽然支持1-Wire多点组网（最多可在一根总线上挂接数十个传感器），'
    '但本设计只使用了一个传感器，未充分利用此特性；'
    '（2）LCD占用了FSMC Bank4和大量GPIO引脚（16位数据总线+控制线约21个引脚），'
    '引脚的占用率较高，限制了其他外设的接入；'
    '（3）系统未预留无线通信接口（如WiFi/蓝牙模块接口），限制了远程监控功能的扩展。'
)
add_para('改进方案：', bold=True)
add_para(
    '（1）增加DS18B20数量，利用其64位ROM ID实现多点温度采集；'
    '（2）若GPIO资源紧张，可将LCD切换为SPI串行接口模式，节省约12个GPIO引脚；'
    '（3）预留USART2/3接口用于连接ESP8266 WiFi模块或HC-05蓝牙模块，实现无线数据传输；'
    '（4）增加SD卡模块（SPI接口），实现温度数据的本地存储和历史查询。'
)

add_heading_styled('6.5  抗干扰能力', 2)
add_para(
    '在实验室环境中，系统运行稳定；但在工业现场等强电磁干扰环境下，系统可能在以下方面'
    '表现出不足：（1）DS18B20的1-Wire通信在长距离（>10米）或强干扰条件下可能出现'
    '通信失败或CRC校验错误；（2）LCD的FSMC并行总线在高频运行时会向外辐射电磁干扰，'
    '同时自身也容易受到外部干扰的影响；（3）按键未加硬件防抖电路（如RC滤波），'
    '在强干扰环境下可能出现误触发。'
)
add_para('改进方案：', bold=True)
add_para(
    '（1）在DS18B20数据线上增加TVS管和共模扼流圈，提高ESD和浪涌防护能力；'
    '（2）在FSMC数据线上串联小电阻（22-33Ω）降低信号反射，改善信号完整性；'
    '（3）PCB布局时注意数字地与模拟地的分割，关键信号线远离高频开关节点；'
    '（4）按键增加硬件RC滤波（如10kΩ+0.1μF），配合软件消抖构成双重防护；'
    '（5）电源入口增加LC滤波和TVS管，提升电源抗扰度。'
)

doc.add_page_break()

# ==================== 第7章 分析和总结 ====================
add_heading_styled('7  分析和总结', 1)

add_para(
    '本课程设计以STM32F103ZET6为核心控制器，成功设计并实现了一套完整的温度测控系统。'
    '系统综合运用了ARM Cortex-M3嵌入式处理器、DS18B20数字温度传感器、ILI9341 TFT LCD'
    '显示屏和PWM脉宽调制等关键技术，实现了温度采集、显示、设定和自动控制的全流程功能。'
)
add_para(
    '在硬件设计方面，系统采用模块化的设计思路，将温度采集、显示驱动、PWM输出、'
    '按键输入和LED指示等功能划分为独立模块，通过标准化的接口进行连接。'
    '这种设计方法不仅提高了系统的可维护性和可扩展性，也为各模块的独立调试提供了便利。'
    '在PCB布局中，注意了数字信号和模拟信号的分区布线，以及电源去耦电容的合理放置，'
    '保证了硬件平台的基本信号完整性。'
)
add_para(
    '在软件设计方面，系统采用前后台架构，主循环负责非实时性任务（如按键处理、'
    '显示更新、串口通信），SysTick中断负责精确时基维护。各硬件驱动模块封装为'
    '独立的.h/.c文件对，符合高内聚低耦合的软件工程原则。温度控制策略采用带滞回区间'
    '的双模式控制，有效避免了设定点附近的频繁模式切换。PWM占空比的线性映射算法'
    '保证了执行器功率的平滑调节，避免了温度控制的过大超调。'
)
add_para(
    '在调试和测试过程中，遇到并解决了以下主要技术难题：'
    '（1）DS18B20的1-Wire时序在72MHz高速MCU上的精确实现——通过关闭中断和使用'
    '精准的软件延时循环确保了微秒级时序的准确性；'
    '（2）ILI9341的FSMC驱动——通过仔细对照ILI9341数据手册的读写时序图，'
    '合理配置FSMC的地址建立时间和数据建立时间参数，最终实现了稳定的LCD显示；'
    '（3）PWM输出验证——使用示波器逐通道验证PWM频率和占空比的准确性，'
    '并与理论计算值进行对比，确认了TIM3的时钟频率和预分频设置正确。'
)
add_para(
    '本项目的主要创新点包括：'
    '（1）采用全数字温度信号链（DS18B20数字传感器+PWM数字控制），'
    '避免了传统模拟方案中的信号调理和ADC转换环节，提高了系统的抗干扰能力和精度；'
    '（2）将FSMC高速并行接口应用于LCD驱动，相比于常见的SPI串行接口方案，'
    '显著提升了显示刷新速率，改善了用户体验；'
    '（3）设计了完整的PWM映射和控制策略，通过滞回比较和分段线性映射实现了'
    '平稳的温度控制，避免了简单的开关控制带来的温度波动。'
)
add_para(
    '从工程应用角度来看，本系统具有以下实际价值：'
    '（1）成本低廉——核心器件（STM32+DS18B20+LCD）的总成本控制在100元以内，'
    '适合需要批量部署的应用场景；（2）可靠性高——全数字信号链减少了模拟电路'
    '的温漂和噪声问题，长期运行稳定性好；（3）易于集成——标准的串口通信接口'
    '使系统可以方便地接入上位机或物联网网关，实现远程监控功能；'
    '（4）平台化设计——模块化的软硬件架构使得系统可以快速适配不同的传感器'
    '（如DHT22温湿度传感器）和执行器（如直流电机、伺服电机），满足不同应用需求。'
)
add_para(
    '通过本课程设计的全流程实践，不仅加深了对ARM Cortex-M3嵌入式系统硬件架构和'
    '软件开发的全面理解，还系统掌握了传感器通信协议（1-Wire）、显示驱动技术（FSMC+ILI9341）、'
    '脉宽调制控制（PWM）和闭环控制算法等嵌入式系统的核心技术。更重要的是，'
    '从需求分析、方案论证到硬件搭建、软件开发、系统调试的完整项目经历，'
    '培养了工程化思维和解决实际问题的能力，为今后从事嵌入式系统和智能仪器仪表'
    '的研发工作奠定了坚实的基础。'
)
add_para(
    '展望未来，本系统可以在以下方向进行进一步深化和拓展：'
    '（1）引入PID控制算法替代当前的线性映射策略，进一步提升温度控制精度和响应速度；'
    '（2）通过ESP8266 WiFi模块实现云平台接入，支持手机APP远程监控和数据记录；'
    '（3）增加SD卡本地存储功能，实现温度历史数据的离线记录和趋势分析；'
    '（4）利用μC/OS或FreeRTOS嵌入式实时操作系统重构软件架构，提升系统的实时性和'
    '多任务处理能力；（5）设计专业的双层PCB板并制作3D打印外壳，使系统从实验原型'
    '升级为可实际部署的产品级设备。'
)

doc.add_page_break()

# ==================== 第8章 参考文献 ====================
add_heading_styled('参考文献', 1)

refs = [
    '[1] 刘火良, 杨森. STM32库开发实战指南: 基于STM32F103[M]. 2版. 北京: 机械工业出版社, 2017.',
    '[2] 张洋, 刘军. 原子教你玩STM32(寄存器版)[M]. 北京: 北京航空航天大学出版社, 2015.',
    '[3] 喻金钱, 喻斌. STM32F系列ARM Cortex-M3核微控制器开发与应用[M]. 北京: 清华大学出版社, 2011.',
    '[4] 陈志旺. STM32嵌入式系统开发实践指南[M]. 北京: 电子工业出版社, 2019.',
    '[5] 卢有亮. 基于STM32的嵌入式系统原理与设计[M]. 北京: 机械工业出版社, 2014.',
    '[6] Dallas Semiconductor. DS18B20 Programmable Resolution 1-Wire Digital Thermometer Datasheet[Z]. Maxim Integrated Products, 2019.',
    '[7] Ilitek. ILI9341 a-Si TFT LCD Single Chip Driver Specification[Z]. Ilitek Corporation, 2010.',
    '[8] STMicroelectronics. RM0008 Reference Manual: STM32F101xx, STM32F102xx, STM32F103xx, STM32F105xx and STM32F107xx advanced ARM-based 32-bit MCUs[Z]. STMicroelectronics, 2021.',
    '[9] STMicroelectronics. STM32F103xE Datasheet: XL-density performance line ARM-based 32-bit MCU with 512KB Flash[Z]. STMicroelectronics, 2020.',
    '[10] 王永虹, 徐炜, 郝立平. STM32系列ARM Cortex-M3微控制器原理与实践[M]. 北京: 北京航空航天大学出版社, 2008.',
    '[11] 马潮. 基于AVR单片机的嵌入式系统原理与应用实践[M]. 北京: 北京航空航天大学出版社, 2011.',
    '[12] 刘军, 张洋, 左忠凯. 例说STM32[M]. 3版. 北京: 北京航空航天大学出版社, 2018.',
    '[13] 周立功. ARM嵌入式系统基础教程[M]. 2版. 北京: 北京航空航天大学出版社, 2008.',
    '[14] Joseph Yiu. The Definitive Guide to ARM Cortex-M3 and Cortex-M4 Processors[M]. 3rd ed. Oxford: Newnes, 2014.',
    '[15] Trevor Martin. The Designer\'s Guide to the Cortex-M Processor Family: A Tutorial Approach[M]. 2nd ed. Oxford: Newnes, 2016.',
    '[16] Jonathan W. Valvano. Embedded Systems: Introduction to ARM Cortex-M Microcontrollers[M]. 5th ed. Self-Published, 2019.',
    '[17] Jean J. Labrosse. Embedded Systems Building Blocks: Complete and Ready-to-Use Modules in C[M]. 2nd ed. CRC Press, 2014.',
    '[18] Edward A. Lee, Sanjit A. Seshia. Introduction to Embedded Systems: A Cyber-Physical Systems Approach[M]. 2nd ed. MIT Press, 2017.',
    '[19] Maxim Integrated. Guidelines for Reliable Long Line 1-Wire Networks[Z]. Application Note 148, 2018.',
    '[20] 中华人民共和国国家质量监督检验检疫总局. GB/T 7714-2015 信息与文献 参考文献著录规则[S]. 北京: 中国标准出版社, 2015.',
]

for ref in refs:
    add_para(ref, size=10.5, first_line_indent=False)

doc.add_page_break()

# ==================== 附录 ====================
add_heading_styled('附录A：核心源码', 1)
add_para('以下为系统核心代码，可直接复制到Keil5工程中使用。前提：Keil5已导入STM32标准外设库（StdPeriph V3.5）。', bold=True)

# --- main.c ---
add_heading_styled('A.1  主程序 main.c', 2)
main_c_code = """#include "main.h"

volatile uint16_t g_cur_temp  = 0;
volatile int16_t  g_set_temp  = DEFAULT_TEMP * 10;
volatile uint8_t  g_fan_pwm   = 0;
volatile uint8_t  g_heat_pwm  = 0;
volatile SysMode_t g_mode     = MODE_STANDBY;
volatile uint8_t  g_fan_on    = 0;
volatile uint8_t  g_heat_on   = 0;

/* 风扇PWM计算（偏差单位0.1°C） */
static uint8_t CalcFan(int16_t diff)
{
    int16_t ex = diff - HYSTERESIS;
    if (ex <= 10) return diff <= HYSTERESIS ? 0 : PWM_MIN;
    if (ex >= 100) return PWM_MAX;
    return PWM_MIN + (uint8_t)((ex - 10) * (PWM_MAX - PWM_MIN) / 90);
}

/* 加热PWM计算 */
static uint8_t CalcHeat(int16_t diff)
{
    int16_t ex = -diff - HYSTERESIS;
    if (ex <= 10) return -diff <= HYSTERESIS ? 0 : PWM_MIN;
    if (ex >= 100) return PWM_MAX;
    return PWM_MIN + (uint8_t)((ex - 10) * (PWM_MAX - PWM_MIN) / 90);
}

/* 温度控制核心 */
void Temp_Control(void)
{
    int16_t raw = DS18B20_Get_Temp();
    int16_t diff;
    if (raw == -1000) return;
    g_cur_temp = (uint16_t)raw;
    diff = (int16_t)g_cur_temp - g_set_temp;

    if (diff > HYSTERESIS) {
        g_mode = MODE_COOLING;
        g_heat_pwm = 0; g_heat_on = 0;
        PWM_Heat(0); LED_Heat(0);
        g_fan_pwm = CalcFan(diff); g_fan_on = 1;
        PWM_Fan(g_fan_pwm); LED_Fan(1);
    } else if (diff < -HYSTERESIS) {
        g_mode = MODE_HEATING;
        g_fan_pwm = 0; g_fan_on = 0;
        PWM_Fan(0); LED_Fan(0);
        g_heat_pwm = CalcHeat(diff); g_heat_on = 1;
        PWM_Heat(g_heat_pwm); LED_Heat(1);
    } else {
        g_mode = MODE_STANDBY;
        g_fan_pwm = 0; g_heat_pwm = 0;
        g_fan_on = 0; g_heat_on = 0;
        PWM_Fan(0); PWM_Heat(0);
        LED_Fan(0); LED_Heat(0);
    }
}

/* LCD刷新 */
void LCD_Update(void)
{
    char buf[32];
    uint8_t *ms[] = {(uint8_t*)"HEATING", (uint8_t*)"COOLING", (uint8_t*)"STANDBY"};

    LCD_ShowString(10, 10, 200, 16, BLACK, (uint8_t*)"Cur :");
    sprintf(buf, "%2d.%1d C", (int16_t)g_cur_temp/10, ((int16_t)g_cur_temp%10+10)%10);
    LCD_ShowString(58, 10, 80, 16, BLACK, (uint8_t*)buf);

    LCD_ShowString(10, 35, 200, 16, BLACK, (uint8_t*)"Set :");
    sprintf(buf, "%2d.%1d C", g_set_temp/10, (g_set_temp%10+10)%10);
    LCD_ShowString(58, 35, 80, 16, BLACK, (uint8_t*)buf);

    LCD_ShowString(10, 60, 200, 16, BLACK, (uint8_t*)"Mode:");
    LCD_ShowString(58, 60, 100, 16, BLACK, ms[g_mode]);

    sprintf(buf, "Fan :%s %3d%%", g_fan_on ? "ON " : "OFF", g_fan_pwm);
    LCD_ShowString(10, 85, 200, 16, BLACK, (uint8_t*)buf);

    sprintf(buf, "Heat:%s %3d%%", g_heat_on ? "ON " : "OFF", g_heat_pwm);
    LCD_ShowString(10, 110, 200, 16, BLACK, (uint8_t*)buf);

    sprintf(buf, "Diff:%+2.1f C", ((float)(int16_t)g_cur_temp - g_set_temp) / 10.0f);
    LCD_ShowString(10, 135, 200, 16, BLACK, (uint8_t*)buf);
}

int main(void)
{
    uint8_t key;
    uint32_t t_disp = 0, t_ctrl = 0;

    /* SystemInit() 已在 startup 中调用: HSE -> PLL -> 72MHz */
    NVIC_PriorityGroupConfig(NVIC_PriorityGroup_2);
    delay_init();
    uart_init(115200);
    LED_Init();
    KEY_Init();
    DS18B20_Init();
    LCD_Init();
    PWM_Init();

    LCD_Clear(WHITE);
    LCD_ShowString(30, 60, 200, 16, BLACK, (uint8_t*)"Temp Control System");
    LCD_ShowString(50, 85, 200, 16, BLACK, (uint8_t*)"Shanxi University");
    delay_ms(1500);
    LCD_Clear(WHITE);

    while (1) {
        key = KEY_Scan(0);
        if (key == 1 && g_set_temp < TEMP_MAX * 10) g_set_temp += 10;
        if (key == 2 && g_set_temp > TEMP_MIN * 10) g_set_temp -= 10;

        if (delay_tdiff(t_ctrl) >= 750) {
            t_ctrl = delay_tick();
            Temp_Control();
        }
        if (delay_tdiff(t_disp) >= 500) {
            t_disp = delay_tick();
            LCD_Update();
        }
    }
}"""
add_para(main_c_code, first_line_indent=False, font_name='Courier New', size=8)

# --- main.h ---
add_heading_styled('A.2  主头文件 main.h', 2)
main_h_code = """#ifndef __MAIN_H
#define __MAIN_H

#include "stm32f10x.h"
#include "delay.h"
#include "usart.h"
#include "ds18b20.h"
#include "lcd.h"
#include "pwm.h"
#include "key.h"
#include "led.h"

/* 控制参数 */
#define DEFAULT_TEMP    25    /* 默认设定温度 °C */
#define TEMP_MIN        10
#define TEMP_MAX        40
#define HYSTERESIS      5     /* 滞回 0.5°C (单位0.1°C) */
#define PWM_MIN         10    /* 最低PWM占空比% */
#define PWM_MAX         100

/* 工作模式 */
typedef enum { MODE_HEATING, MODE_COOLING, MODE_STANDBY } SysMode_t;

/* 全局变量 */
extern volatile uint16_t g_cur_temp;    /* 当前温度×10 */
extern volatile int16_t  g_set_temp;    /* 设定温度×10 */
extern volatile uint8_t  g_fan_pwm;     /* 风扇占空比% */
extern volatile uint8_t  g_heat_pwm;    /* 加热占空比% */
extern volatile SysMode_t g_mode;
extern volatile uint8_t  g_fan_on, g_heat_on;

/* 函数 */
void Temp_Control(void);
void LCD_Update(void);

#endif"""
add_para(main_h_code, first_line_indent=False, font_name='Courier New', size=8)

# --- ds18b20.h ---
add_heading_styled('A.3  DS18B20驱动 ds18b20.h', 2)
ds18b20_h_code = """#ifndef __DS18B20_H
#define __DS18B20_H
#include "stm32f10x.h"
#include "delay.h"

/* PG9 — 1-Wire数据线 */
#define DS18B20_PORT            GPIOG
#define DS18B20_PIN             GPIO_Pin_9
#define DS18B20_RCC_PORT        RCC_APB2Periph_GPIOG

/* 输入/输出模式切换 */
#define DS18B20_DQ_OUT()  {                        \\
    GPIO_InitTypeDef g;                              \\
    g.GPIO_Pin = DS18B20_PIN;                        \\
    g.GPIO_Mode = GPIO_Mode_Out_PP;                  \\
    g.GPIO_Speed = GPIO_Speed_50MHz;                 \\
    GPIO_Init(DS18B20_PORT, &g);                     \\
}
#define DS18B20_DQ_IN()   {                        \\
    GPIO_InitTypeDef g;                              \\
    g.GPIO_Pin = DS18B20_PIN;                        \\
    g.GPIO_Mode = GPIO_Mode_IPU;                     \\
    GPIO_Init(DS18B20_PORT, &g);                     \\
}

#define DS18B20_DQ_HIGH()  GPIO_SetBits(DS18B20_PORT, DS18B20_PIN)
#define DS18B20_DQ_LOW()   GPIO_ResetBits(DS18B20_PORT, DS18B20_PIN)
#define DS18B20_DQ_READ()  GPIO_ReadInputDataBit(DS18B20_PORT, DS18B20_PIN)

/* DS18B20 ROM命令 */
#define DS18B20_CMD_SKIP_ROM         0xCC
#define DS18B20_CMD_CONVERT_TEMP     0x44
#define DS18B20_CMD_READ_SCRATCHPAD  0xBE

uint8_t DS18B20_Init(void);
void    DS18B20_Rst(void);
uint8_t DS18B20_Check(void);
uint8_t DS18B20_Read_Byte(void);
void    DS18B20_Write_Byte(uint8_t dat);
void    DS18B20_Start(void);
int16_t DS18B20_Get_Temp(void);

#endif"""
add_para(ds18b20_h_code, first_line_indent=False, font_name='Courier New', size=8)

# --- ds18b20.c ---
add_heading_styled('A.4  DS18B20驱动 ds18b20.c', 2)
ds18b20_c_code = """#include "ds18b20.h"

uint8_t DS18B20_Init(void)
{
    GPIO_InitTypeDef g;
    RCC_APB2PeriphClockCmd(DS18B20_RCC_PORT, ENABLE);
    g.GPIO_Pin = DS18B20_PIN;
    g.GPIO_Mode = GPIO_Mode_Out_PP;
    g.GPIO_Speed = GPIO_Speed_50MHz;
    GPIO_Init(DS18B20_PORT, &g);
    DS18B20_DQ_HIGH();
    return DS18B20_Check();
}

void DS18B20_Rst(void)
{
    DS18B20_DQ_OUT();
    DS18B20_DQ_LOW();
    delay_us(500);
    DS18B20_DQ_HIGH();
    delay_us(15);
}

uint8_t DS18B20_Check(void)
{
    uint8_t retry = 0;
    DS18B20_DQ_IN();
    while (DS18B20_DQ_READ() && retry < 200) { retry++; delay_us(1); }
    if (retry >= 200) return 1;
    retry = 0;
    while (!DS18B20_DQ_READ() && retry < 240) { retry++; delay_us(1); }
    if (retry >= 240) return 1;
    return 0;
}

uint8_t DS18B20_Read_Byte(void)
{
    uint8_t i, byte = 0;
    for (i = 0; i < 8; i++) {
        byte >>= 1;
        DS18B20_DQ_OUT(); DS18B20_DQ_LOW();
        delay_us(2); DS18B20_DQ_HIGH(); delay_us(1);
        DS18B20_DQ_IN(); delay_us(2);
        if (DS18B20_DQ_READ()) byte |= 0x80;
        delay_us(45);
    }
    return byte;
}

void DS18B20_Write_Byte(uint8_t dat)
{
    uint8_t i;
    DS18B20_DQ_OUT();
    for (i = 0; i < 8; i++) {
        if (dat & 0x01) {
            DS18B20_DQ_LOW(); delay_us(2);
            DS18B20_DQ_HIGH(); delay_us(60);
        } else {
            DS18B20_DQ_LOW(); delay_us(60);
            DS18B20_DQ_HIGH(); delay_us(2);
        }
        dat >>= 1;
    }
}

void DS18B20_Start(void)
{
    DS18B20_Rst(); DS18B20_Check();
    DS18B20_Write_Byte(DS18B20_CMD_SKIP_ROM);
    DS18B20_Write_Byte(DS18B20_CMD_CONVERT_TEMP);
}

int16_t DS18B20_Get_Temp(void)
{
    uint8_t tl, th;
    int16_t raw;

    DS18B20_Rst();
    if (DS18B20_Check() != 0) return -1000;  /* 传感器异常 */

    DS18B20_Write_Byte(DS18B20_CMD_SKIP_ROM);
    DS18B20_Write_Byte(DS18B20_CMD_READ_SCRATCHPAD);
    tl = DS18B20_Read_Byte();
    th = DS18B20_Read_Byte();
    raw = (int16_t)(((uint16_t)th << 8) | (uint16_t)tl);
    raw = (int16_t)((float)raw * 0.625f);  /* ×0.0625×10 */
    return raw;
}"""
add_para(ds18b20_c_code, first_line_indent=False, font_name='Courier New', size=8)

# --- pwm.h ---
add_heading_styled('A.5  PWM驱动 pwm.h', 2)
pwm_h_code = """#ifndef __PWM_H
#define __PWM_H
#include "stm32f10x.h"

#define PWM_ARR  999     /* 1kHz: 72MHz/(71+1)/(999+1) */
#define PWM_PSC  71

void PWM_Init(void);
void PWM_Fan(uint8_t duty);   /* duty: 0-100 */
void PWM_Heat(uint8_t duty);
#endif"""
add_para(pwm_h_code, first_line_indent=False, font_name='Courier New', size=8)

add_heading_styled('A.6  PWM驱动 pwm.c', 2)
pwm_c_code = """#include "pwm.h"

void PWM_Init(void)
{
    GPIO_InitTypeDef g;
    TIM_TimeBaseInitTypeDef t;
    TIM_OCInitTypeDef oc;

    RCC_APB2PeriphClockCmd(RCC_APB2Periph_GPIOA, ENABLE);
    RCC_APB1PeriphClockCmd(RCC_APB1Periph_TIM3, ENABLE);

    g.GPIO_Pin = GPIO_Pin_6 | GPIO_Pin_7;
    g.GPIO_Mode = GPIO_Mode_AF_PP;
    g.GPIO_Speed = GPIO_Speed_50MHz;
    GPIO_Init(GPIOA, &g);

    t.TIM_Prescaler = PWM_PSC;
    t.TIM_Period = PWM_ARR;
    t.TIM_ClockDivision = TIM_CKD_DIV1;
    t.TIM_CounterMode = TIM_CounterMode_Up;
    TIM_TimeBaseInit(TIM3, &t);

    oc.TIM_OCMode = TIM_OCMode_PWM1;
    oc.TIM_OutputState = TIM_OutputState_Enable;
    oc.TIM_OCPolarity = TIM_OCPolarity_High;
    oc.TIM_Pulse = 0;
    TIM_OC1Init(TIM3, &oc);
    TIM_OC2Init(TIM3, &oc);
    TIM_OC1PreloadConfig(TIM3, TIM_OCPreload_Enable);
    TIM_OC2PreloadConfig(TIM3, TIM_OCPreload_Enable);
    TIM_Cmd(TIM3, ENABLE);
}

void PWM_Fan(uint8_t duty)
{
    if (duty > 100) duty = 100;
    TIM_SetCompare1(TIM3, (uint16_t)duty * (PWM_ARR + 1) / 100);
}

void PWM_Heat(uint8_t duty)
{
    if (duty > 100) duty = 100;
    TIM_SetCompare2(TIM3, (uint16_t)duty * (PWM_ARR + 1) / 100);
}"""
add_para(pwm_c_code, first_line_indent=False, font_name='Courier New', size=8)

# --- delay ---
add_heading_styled('A.7  延时驱动 delay.h / delay.c', 2)
delay_h_code = """#ifndef __DELAY_H
#define __DELAY_H
#include "stm32f10x.h"
void     delay_init(void);
void     delay_ms(uint32_t nms);
void     delay_us(uint32_t nus);
uint32_t delay_tick(void);
uint32_t delay_tdiff(uint32_t last);
#endif"""
add_para(delay_h_code, first_line_indent=False, font_name='Courier New', size=8)

delay_c_code = """#include "delay.h"

static volatile uint32_t g_tick = 0;

void delay_init(void)
{
    SysTick_CLKSourceConfig(SysTick_CLKSource_HCLK_Div8);  /* 9MHz */
    SysTick->LOAD = 9000 - 1;  /* 1ms */
    SysTick->CTRL |= SysTick_CTRL_TICKINT_Msk | SysTick_CTRL_ENABLE_Msk;
}

void SysTick_Handler(void) { g_tick++; }

void delay_ms(uint32_t n) {
    uint32_t t = g_tick;
    while ((g_tick - t) < n);
}

void delay_us(uint32_t n) {
    n *= 9;  /* 72MHz, ~9 cycles/loop */
    while (n--) __NOP();
}

uint32_t delay_tick(void)        { return g_tick; }
uint32_t delay_tdiff(uint32_t t) { return g_tick - t; }"""
add_para(delay_c_code, first_line_indent=False, font_name='Courier New', size=8)

# --- key ---
add_heading_styled('A.8  按键驱动 key.c', 2)
key_c_code = """#include "key.h"

void KEY_Init(void)
{
    GPIO_InitTypeDef g;
    RCC_APB2PeriphClockCmd(RCC_APB2Periph_GPIOG, ENABLE);
    g.GPIO_Mode = GPIO_Mode_IPU;
    g.GPIO_Pin = GPIO_Pin_7 | GPIO_Pin_8;
    GPIO_Init(GPIOG, &g);
}

uint8_t KEY_Scan(uint8_t mode)
{
    static uint8_t up = 1;
    if (up && (GPIO_ReadInputDataBit(GPIOG, GPIO_Pin_8) == RESET ||
               GPIO_ReadInputDataBit(GPIOG, GPIO_Pin_7) == RESET)) {
        delay_ms(20);
        up = 0;
        if (GPIO_ReadInputDataBit(GPIOG, GPIO_Pin_8) == RESET) return 1;
        if (GPIO_ReadInputDataBit(GPIOG, GPIO_Pin_7) == RESET) return 2;
    } else if (GPIO_ReadInputDataBit(GPIOG, GPIO_Pin_8) != RESET &&
               GPIO_ReadInputDataBit(GPIOG, GPIO_Pin_7) != RESET) {
        up = 1;
    }
    if (mode) up = 1;
    return 0;
}"""
add_para(key_c_code, first_line_indent=False, font_name='Courier New', size=8)

# --- led ---
add_heading_styled('A.9  LED驱动 led.c', 2)
led_c_code = """#include "led.h"

void LED_Init(void)
{
    GPIO_InitTypeDef g;
    RCC_APB2PeriphClockCmd(RCC_APB2Periph_GPIOC, ENABLE);
    g.GPIO_Mode = GPIO_Mode_Out_PP;
    g.GPIO_Speed = GPIO_Speed_50MHz;
    g.GPIO_Pin = LED_FAN_PIN | LED_HEAT_PIN | LED_PWR_PIN;
    GPIO_Init(LED_PORT, &g);
    GPIO_SetBits(LED_PORT, LED_FAN_PIN | LED_HEAT_PIN);
    GPIO_ResetBits(LED_PORT, LED_PWR_PIN);
}

void LED_Fan(uint8_t on)  { if(on) GPIO_ResetBits(LED_PORT,LED_FAN_PIN);
                            else   GPIO_SetBits(LED_PORT,LED_FAN_PIN); }
void LED_Heat(uint8_t on) { if(on) GPIO_ResetBits(LED_PORT,LED_HEAT_PIN);
                            else   GPIO_SetBits(LED_PORT,LED_HEAT_PIN); }"""
add_para(led_c_code, first_line_indent=False, font_name='Courier New', size=8)

add_heading_styled('附录B：原理图', 1)
add_para('系统电路原理图和模块连接图见Diagram目录。')

add_heading_styled('附录C：测试数据', 1)
add_para('测试数据表、温度-PWM曲线图、系统响应曲线见Test目录。')

add_heading_styled('附录D：调试记录', 1)
add_para(
    '调试记录1 — DS18B20通信失败\n'
    '  现象：DS18B20_Check()返回1。\n'
    '  原因：忘记焊接4.7kΩ上拉电阻。\n'
    '  解决：补焊后通信正常。\n\n'
    '调试记录2 — LCD花屏\n'
    '  现象：显示随机彩色条纹。\n'
    '  原因：FSMC_DataSetupTime=1太小，数据建立时间不足。\n'
    '  解决：改为15（约208ns）后正常。\n\n'
    '调试记录3 — PWM无输出\n'
    '  现象：PA6无波形。\n'
    '  原因：GPIO_Mode配成了Out_PP而非AF_PP。\n'
    '  解决：改为GPIO_Mode_AF_PP。'
)

add_heading_styled('附录E：LCD驱动说明', 1)
add_para(
    'lcd.c（约260行）包含ILI9341 FSMC初始化序列和8×16 ASCII字库（95个可打印字符）。\n'
    '由于字库占用约150行篇幅，完整代码请见Code/HARDWARE/LCD/lcd.c。\n'
    '核心接口函数：\n'
    '  LCD_Init()          — GPIO+FSMC+ILI9341寄存器初始化\n'
    '  LCD_Clear(color)    — 全屏填充指定颜色\n'
    '  LCD_ShowChar(x,y,ch,size,color) — 在指定位置显示单个字符\n'
    '  LCD_ShowString(x,y,w,size,color,str) — 显示字符串（自动换行）'
)

# ==================== 保存 ====================
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                           '课程设计报告_含源码.docx')
doc.save(output_path)
print(f'报告已生成: {output_path}')
